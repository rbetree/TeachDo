import asyncio
import json
import io
import random
import re
import os
import sys
from pathlib import Path
from fastapi import FastAPI, UploadFile, File
import time
import logging
from pydantic import BaseModel
import uuid
import httpx
from urllib.parse import quote, urlsplit
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi import UploadFile, File, HTTPException, Form
from fastapi import FastAPI, HTTPException, Query, Request, Response
from typing import AsyncGenerator, Literal, Any
from starlette.background import BackgroundTask
try:
    # 兼容在 `backend/main_api` 目录下直接运行（例如 `uvicorn main:app`）
    from outline_client import A2AOutlineClientWrapper
    from content_client import A2AContentClientWrapper
except ImportError:  # pragma: no cover - 兼容以包方式导入（用于单元测试等）
    from backend.main_api.outline_client import A2AOutlineClientWrapper
    from backend.main_api.content_client import A2AContentClientWrapper

logger = logging.getLogger(__name__)

def _find_repo_root(start: Path) -> Path:
    """
    向上查找项目根目录：
    - 优先命中 `.git/` 或 `env_template.txt`
    - 若不存在（例如 docker 镜像只拷贝了单服务目录），退化为包含 `main.py` 的目录
    """
    start_dir = start if start.is_dir() else start.parent
    fallback_service_root: Path | None = None

    current = start_dir
    while True:
        if (current / ".git").exists() or (current / "env_template.txt").exists():
            return current
        if fallback_service_root is None and (current / "main.py").exists():
            fallback_service_root = current

        parent = current.parent
        if parent == current:
            break
        current = parent

    return fallback_service_root or Path.cwd()


_repo_root = _find_repo_root(Path(__file__).resolve())
# 允许在 `backend/main_api` 目录下直接运行（例如 `python main.py`）：确保可导入 backend.*
if str(_repo_root) not in sys.path:
    sys.path.insert(0, str(_repo_root))

from backend.common.env_loader import load_env_files
from backend.common.course_outputs_injection import (
    COURSE_OUTPUTS_START_MARKER,
    build_course_outputs_injection_markdown,
)
from backend.common.cors import get_cors_middleware_kwargs
from backend.common.proxy_guard import (
    get_proxy_allowed_hosts,
    get_proxy_max_bytes,
    is_proxy_host_allowed,
)
from backend.common.static_files import resolve_safe_static_file
from backend.common.settings_store import access_host_for_bind_host
from backend.common.url_security import (
    REDIRECT_STATUS_CODES,
    UrlAccessError,
    resolve_and_validate_redirect_url,
    validate_public_http_url,
)

load_env_files(repo_root=_repo_root, service_dir=Path(__file__).resolve().parent)

TEMPLATE_DIR = Path(__file__).resolve().parent / "template"


def _get_outline_api() -> str:
    return os.environ.get(
        "OUTLINE_API",
        f"http://{access_host_for_bind_host(os.environ.get('HOST', '127.0.0.1'))}:{os.environ.get('OUTLINE_API_PORT', '10001')}",
    )


def _get_content_api() -> str:
    return os.environ.get(
        "CONTENT_API",
        f"http://{access_host_for_bind_host(os.environ.get('HOST', '127.0.0.1'))}:{os.environ.get('CONTENT_API_PORT', '10011')}",
    )


async def _aclose_httpx_stream(response: httpx.Response, client: httpx.AsyncClient) -> None:
    await response.aclose()
    await client.aclose()


async def _open_validated_proxy_stream(
    target_url: str,
    *,
    headers: dict[str, str],
    max_redirects: int = 3,
) -> tuple[httpx.AsyncClient, httpx.Response]:
    """
    以“每次跳转都校验目标地址”的方式打开上游流，避免 SSRF 通过重定向绕过。
    调用方负责在返回后关闭 response/client。
    """
    current_url = validate_public_http_url(target_url)
    allowed_hosts = get_proxy_allowed_hosts()

    def _ensure_allowed(url: str) -> None:
        if not allowed_hosts:
            return
        host = urlsplit(url).hostname or ""
        if not is_proxy_host_allowed(host, allowed_hosts):
            raise UrlAccessError("目标域名不在允许列表", status_code=403)

    _ensure_allowed(current_url)
    client = httpx.AsyncClient(timeout=httpx.Timeout(30.0), follow_redirects=False)

    try:
        for _ in range(max_redirects + 1):
            request = client.build_request("GET", current_url, headers=headers)
            response = await client.send(request, stream=True)
            if response.status_code in REDIRECT_STATUS_CODES and "location" in response.headers:
                next_url = resolve_and_validate_redirect_url(current_url, response.headers["location"])
                _ensure_allowed(next_url)
                await response.aclose()
                current_url = next_url
                continue
            return client, response
    except Exception:
        await client.aclose()
        raise

    await client.aclose()
    raise HTTPException(status_code=502, detail="上游重定向次数过多")
app = FastAPI()

# settings API（允许在前端“设置”页写入 var/settings.json）
try:
    from backend.main_api.settings_api import register_settings_routes

    register_settings_routes(app)
except Exception:  # pragma: no cover - 单服务打包/裁剪场景允许缺失
    pass

# Allow CORS for the frontend development server
app.add_middleware(
    CORSMiddleware,
    **get_cors_middleware_kwargs(allow_credentials=True),
)

class AipptRequest(BaseModel):
    content: str
    language: str
    model: str
    stream: bool


class AssistantChatMessage(BaseModel):
    role: Literal["user", "assistant"]
    content: str


class AssistantMaterialContext(BaseModel):
    title: str
    subject: str | None = None
    description: str | None = None
    objectives: str | None = None


class AssistantChatRequest(BaseModel):
    """
    助教对话请求（全局单会话，历史由前端维护并在每次请求时透传）。
    """

    messages: list[AssistantChatMessage]
    user_id: str = "default_user"
    kb_file_ids: list[str] | None = None
    material: AssistantMaterialContext | None = None
    language: str = "zh"


class LessonPlanProcedureStep(BaseModel):
    step: str
    duration: str
    activity: str


class LessonPlan(BaseModel):
    """
    LessonPlan（与 frontend/types.ts 对齐）
    """

    title: str
    targetAudience: str
    duration: str
    objectives: list[str]
    materials: list[str]
    procedure: list[LessonPlanProcedureStep]
    homework: str


class LessonPlanRequest(BaseModel):
    title: str
    subject: str | None = None
    description: str | None = None
    objectives: str | None = None
    outlineContent: str
    language: str = "zh"
    sessionId: str | None = None
    user_id: str | None = None
    kb_file_ids: list[str] | None = None
    # 教案内容模板（用于影响生成策略/结构，保持与前端 camelCase 一致）
    templateId: str | None = None


class LessonStyle(BaseModel):
    """
    Lesson 导出样式（V1）
    - 作为“展示/导出层”参数，不参与 LessonPlan 内容生成
    """

    fontZh: str = "微软雅黑"
    titleSizePt: int = 20
    h1SizePt: int = 16
    h2SizePt: int = 14
    bodySizePt: int = 12
    lineSpacing: float = 1.5

    # 页边距（cm）
    marginTopCm: float = 2.54
    marginBottomCm: float = 2.54
    marginLeftCm: float = 2.54
    marginRightCm: float = 2.54


class LessonExportDocxRequest(BaseModel):
    lessonPlan: LessonPlan
    style: LessonStyle | None = None
    language: str | None = None
    # 教案 docx 导出模板（与前端保持 camelCase）
    templateId: str | None = None
    # 可选：导出文件持久化到 artifacts（与前端保持 camelCase）
    userId: str | None = None
    materialId: str | None = None
    persist: bool | None = None


def _normalize_lesson_template_id(value: str | None) -> str:
    """
    统一归一化教案模板 ID（生成/预览/导出共用）。
    - 未指定：默认 lesson_simple
    - 兼容常见别名：simple/table/form/jnu 等
    """
    tpl = (value or "").strip() or "lesson_simple"
    if tpl in {"default", "simple"}:
        return "lesson_simple"
    if tpl in {"table"}:
        return "lesson_table"
    if tpl in {"form", "jnu", "jnu_form", "lesson_form"}:
        return "lesson_jnu_form"
    return tpl


def _split_objectives_text(text: str) -> list[str]:
    """
    将自由文本教学目标拆成条目列表（用于无 LLM 的兜底生成）。
    """
    raw = (text or "").strip()
    if not raw:
        return []
    bits = re.split(r"[；;。\n]+", raw)
    cleaned: list[str] = []
    for b in bits:
        item = (b or "").strip()
        if not item:
            continue
        if item in cleaned:
            continue
        cleaned.append(item)
    return cleaned


def _guess_procedure_steps_from_outline(outline_md: str, *, max_steps: int = 6) -> list[str]:
    """
    从 Markdown 大纲里猜测教学流程步骤（用于无 LLM 的兜底生成）。
    """
    steps: list[str] = []
    for raw_line in (outline_md or "").splitlines():
        line = (raw_line or "").strip()
        if not line:
            continue
        if line.startswith("#"):
            title = line.lstrip("#").strip()
            if title and title not in steps:
                steps.append(title)
        elif line.startswith(("-", "*", "+")):
            title = line.lstrip("-*+").strip()
            if title and title not in steps:
                steps.append(title)
        if len(steps) >= max_steps:
            break
    return steps


def _fallback_generate_lesson_plan(req: LessonPlanRequest) -> LessonPlan:
    """
    无 LLM 配置时的兜底教案生成：
    - 保证端到端链路可用（SSE 预览 + docx 导出）
    - 质量不追求最优，但内容结构完整、可读
    """
    lang = (req.language or "zh").strip().lower()
    want_english = lang in {"en", "english"}
    tpl = _normalize_lesson_template_id(req.templateId)
    if tpl not in {"lesson_simple", "lesson_table", "lesson_jnu_form"}:
        tpl = "lesson_simple"

    title = (req.title or "").strip() or ("Lesson Plan" if want_english else "教案")
    objectives = _split_objectives_text(req.objectives or "")
    if not objectives:
        objectives = [
            "理解本节课核心概念与关键结论" if not want_english else "Understand the key concepts and core conclusions",
            "能完成基础例题/练习并进行简单迁移" if not want_english else "Solve basic exercises and apply the concept",
        ]
    if tpl == "lesson_jnu_form":
        # 表单模板：目标用于“重点/难点”，确保最后一条为“难点：...”
        diff = None
        rest: list[str] = []
        for item in objectives:
            text = str(item or "").strip()
            if not text:
                continue
            is_diff = bool(re.match(r"^difficulty\s*[:：]", text, re.IGNORECASE)) if want_english else bool(re.match(r"^难点\s*[:：]", text))
            if is_diff and diff is None:
                diff = text
                continue
            rest.append(text)

        if diff is None:
            topic = (req.title or "").strip() or ("this lesson" if want_english else "本节课")
            diff = (
                f"Difficulty: Explain and apply the key reasoning of {topic} in new contexts."
                if want_english
                else f"难点：掌握{topic}的关键推导与应用，并能迁移解决典型问题。"
            )

        objectives = rest + [diff]

    steps = _guess_procedure_steps_from_outline(req.outlineContent or "", max_steps=6)
    if not steps:
        steps = [
            "导入与复习" if not want_english else "Warm-up & review",
            "新知讲解" if not want_english else "Concept introduction",
            "例题与练习" if not want_english else "Examples & practice",
            "总结提升" if not want_english else "Wrap-up",
        ]

    # 简单分配时长（总时长默认 45 分钟）
    total_minutes = 45
    minutes_each = max(5, int(round(total_minutes / max(1, len(steps)))))
    procedure: list[LessonPlanProcedureStep] = []
    for idx, name in enumerate(steps):
        m = minutes_each
        duration = f"{m}分钟" if not want_english else f"{m} min"
        activity = (
            f"围绕「{name}」组织讲解与互动，包含提问、板书要点与即时练习。"
            if not want_english
            else f"Teach and interact around “{name}”, including questions, key points, and quick practice."
        )
        procedure.append(
            LessonPlanProcedureStep(
                step=f"{idx + 1}. {name}",
                duration=duration,
                activity=activity,
            )
        )

    materials = [
        "课件/PPT",
        "板书/白板",
        "练习题/作业纸",
    ] if not want_english else ["Slides", "Whiteboard", "Practice sheets"]
    if tpl == "lesson_jnu_form":
        # 参考资料占位（便于表单模板直接落地填写）
        materials = list(materials) + (["教材/参考书：______"] if not want_english else ["Textbook/Reference: ______"])

    homework = (
        "完成课后练习 1~3 题，并用自己的话总结本课关键结论（不少于 100 字）。"
        if not want_english
        else "Finish exercises 1–3 and summarize the key takeaway in your own words (100+ words)."
    )

    return LessonPlan(
        title=title,
        targetAudience=("中学学生" if not want_english else "Students"),
        duration=("45分钟" if not want_english else "45 min"),
        objectives=objectives,
        materials=materials,
        procedure=procedure,
        homework=homework,
    )


def _try_get_lesson_llm_settings() -> dict[str, str] | None:
    """
    Lesson 生成的模型配置：
    - 优先读取 LESSON_*，未配置则回退到 OUTLINE_*（减少新增配置成本）
    - 仅支持 openai 兼容协议（base_url + /chat/completions）
    - 若缺少必要配置，返回 None（由兜底生成保证链路可用）
    """
    llm_type = (os.getenv("LESSON_TYPE") or os.getenv("OUTLINE_TYPE") or "").strip().lower()
    llm_model = (os.getenv("LESSON_MODEL") or os.getenv("OUTLINE_MODEL") or "").strip()
    llm_api_key = (os.getenv("LESSON_API_KEY") or os.getenv("OUTLINE_API_KEY") or "").strip()
    llm_base_url = (os.getenv("LESSON_BASE_URL") or os.getenv("OUTLINE_BASE_URL") or "https://api.openai.com/v1").strip()

    if not llm_type or not llm_model or not llm_api_key:
        return None
    openai_compatible = {"openai", "ollama", "vllm", "local_openai", "xinference"}
    if llm_type not in openai_compatible:
        raise RuntimeError(f"当前 Lesson 仅支持 openai 兼容协议，检测到 LESSON_TYPE/OUTLINE_TYPE={llm_type}")
    return {"type": llm_type, "model": llm_model, "api_key": llm_api_key, "base_url": llm_base_url}


def _strip_json_code_fence(text: str) -> str:
    """
    兼容模型把 JSON 包在 ```json ... ``` 围栏内的情况。
    """
    s = (text or "").strip()
    if not s.startswith("```"):
        return s
    lines = s.splitlines()
    if len(lines) < 2:
        return s
    if not lines[0].strip().startswith("```"):
        return s
    end_idx = -1
    for i in range(len(lines) - 1, -1, -1):
        if lines[i].strip() == "```":
            end_idx = i
            break
    if end_idx <= 0:
        return s
    return "\n".join(lines[1:end_idx]).strip()


def _extract_first_json_object(text: str) -> dict[str, Any]:
    """
    从模型输出中提取第一个 JSON 对象。
    - 允许前后夹杂说明文字
    - 失败则抛出异常，由上层决定兜底策略
    """
    s = _strip_json_code_fence(text)
    start = s.find("{")
    end = s.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("未找到 JSON 对象边界")
    candidate = s[start : end + 1].strip()
    obj = json.loads(candidate)
    if not isinstance(obj, dict):
        raise ValueError("JSON 顶层不是对象")
    return obj


def _build_lesson_system_prompt(*, req: LessonPlanRequest, full_context: str, kb_context: str) -> str:
    """
    Lesson 生成 system prompt（用于 LLM 路径）。
    """
    lang = (req.language or "zh").strip().lower()
    want_english = lang in {"en", "english"}
    tpl = _normalize_lesson_template_id(req.templateId)
    if tpl not in {"lesson_simple", "lesson_table", "lesson_jnu_form"}:
        tpl = "lesson_simple"

    if want_english:
        base = (
            "You are TeachDo's lesson plan generator.\n"
            "Generate a structured lesson plan based on the provided outline.\n"
            "Rules:\n"
            "- Follow the outline structure and do NOT invent topics that are unrelated.\n"
            "- If course outputs (full text, not retrieved) are provided, treat them as reference-only for consistency; do NOT copy verbatim and do NOT reuse its structure.\n"
            "- Output must be STRICT JSON only (no markdown, no code fences).\n"
            "- Keep it practical for classroom use.\n"
        )
    else:
        base = (
            "你是 TeachDo 的教案生成器。\n"
            "请基于给定的大纲生成结构化教案。\n"
            "规则：\n"
            "- 必须参考大纲结构，不要引入无关主题。\n"
            "- 若提供“课程产出（全文，不经检索）”，仅作参考用于术语/事实对齐：不要原文照抄、不要套用其目录/结构，仍以本次大纲为准。\n"
            "- 输出必须是严格 JSON（不要 markdown、不要代码块围栏）。\n"
            "- 内容要可落地、可直接用于课堂。\n"
        )

    if tpl == "lesson_jnu_form":
        if want_english:
            base += (
                "\nTemplate: form-style (fields).\n"
                "- Objectives will be used as Key Points and Difficulty.\n"
                "- Ensure the final objective starts with 'Difficulty:'.\n"
            )
        else:
            base += (
                "\n当前模板：教案表单（字段）。\n"
                "- objectives 将用于“重点/难点”。请确保最后一条以“难点：”开头。\n"
            )

    context_bits: list[str] = []
    title = (req.title or "").strip()
    if title:
        context_bits.append(("Topic: " if want_english else "主题：") + title)
    if req.subject:
        context_bits.append(("Subject: " if want_english else "学科：") + str(req.subject).strip())
    if req.description:
        context_bits.append(("Background: " if want_english else "背景：") + str(req.description).strip())
    if req.objectives:
        context_bits.append(("User objectives: " if want_english else "用户提供的教学目标：") + str(req.objectives).strip())

    outline = (req.outlineContent or "").strip()
    context_bits.append(("Outline (Markdown):\n" if want_english else "课程大纲（Markdown）：\n") + outline)

    if full_context and full_context.strip():
        context_bits.append(
            (
                "Course outputs (full text, not retrieved):\n(Note: reference-only; do NOT copy verbatim; do NOT reuse its structure.)\n"
                if want_english
                else "课程产出（全文，不经检索）：\n（说明：仅供参考，不要原文照抄，不要套用其目录/结构。）\n"
            )
            + full_context.strip()
        )

    if kb_context and kb_context.strip():
        context_bits.append(("Reference snippets (RAG):\n" if want_english else "参考资料检索片段（RAG）：\n") + kb_context.strip())

    return base + "\n\n" + "\n".join(context_bits).strip()


async def _generate_lesson_section_with_llm(
    *,
    llm_settings: dict[str, str],
    system_prompt: str,
    user_prompt: str,
    temperature: float = 0.4,
    max_output_chars: int = 20_000,
) -> dict[str, Any]:
    """
    调用 OpenAI 兼容协议生成一个 JSON 片段（section/meta）。
    """
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]

    text = ""
    async for delta in iter_assistant_text_chunks(
        model=llm_settings["model"],
        api_key=llm_settings["api_key"],
        base_url=llm_settings["base_url"],
        messages=messages,
        temperature=temperature,
    ):
        text += delta
        if len(text) > max_output_chars:
            break

    return _extract_first_json_object(text)


async def stream_lesson_plan_sse(req: LessonPlanRequest) -> AsyncGenerator[bytes, None]:
    """
    LessonPlan SSE：
    - data: {"type":"section",...}
    - data: {"type":"final","data":{LessonPlan}}
    - data: {"type":"error","text":"..."}
    - data: [DONE]
    """
    outline = (req.outlineContent or "").strip()
    if not outline:
        payload = json.dumps({"type": "error", "text": "outlineContent 不能为空"}, ensure_ascii=False, separators=(",", ":"))
        yield _encode_sse_data(payload)
        yield b"data: [DONE]\n\n"
        return

    user_id = str(req.user_id or "default_user")
    resolved_kb_file_ids = _normalize_kb_file_ids(req.kb_file_ids)
    personaldb_url = _get_personaldb_url()

    full_context = ""
    kb_context = ""
    if resolved_kb_file_ids and personaldb_url and await _is_personaldb_ready(personaldb_url):
        # query 选用 “标题 + 大纲” 的组合，尽量贴近教案生成语义
        query = f"{(req.title or '').strip()}\n{outline}".strip()
        full_context, kb_context = await _build_personaldb_kb_contexts(
            personaldb_url,
            user_id=user_id,
            query=query,
            kb_file_ids=resolved_kb_file_ids,
            rag_topk=5,
        )

    llm_settings: dict[str, str] | None = None
    try:
        llm_settings = _try_get_lesson_llm_settings()
    except Exception as exc:
        # 配置异常：作为错误事件返回，但仍保证 [DONE] 收尾
        payload = json.dumps({"type": "error", "text": f"Lesson 模型配置错误：{exc}"}, ensure_ascii=False, separators=(",", ":"))
        yield _encode_sse_data(payload)
        yield b"data: [DONE]\n\n"
        return

    # 兜底：无 LLM 配置时，直接按 outline 生成一个可用版本
    if not llm_settings:
        plan = _fallback_generate_lesson_plan(req)
        yield _encode_sse_data(json.dumps({"type": "section", "section": "objectives", "data": plan.objectives}, ensure_ascii=False, separators=(",", ":")))
        yield _encode_sse_data(json.dumps({"type": "section", "section": "materials", "data": plan.materials}, ensure_ascii=False, separators=(",", ":")))
        yield _encode_sse_data(json.dumps({"type": "section", "section": "procedure", "data": [p.model_dump() for p in plan.procedure]}, ensure_ascii=False, separators=(",", ":")))
        yield _encode_sse_data(json.dumps({"type": "section", "section": "homework", "data": plan.homework}, ensure_ascii=False, separators=(",", ":")))
        yield _encode_sse_data(json.dumps({"type": "final", "data": plan.model_dump()}, ensure_ascii=False, separators=(",", ":")))
        yield b"data: [DONE]\n\n"
        return

    lang = (req.language or "zh").strip().lower()
    want_english = lang in {"en", "english"}
    tpl = _normalize_lesson_template_id(req.templateId)
    if tpl not in {"lesson_simple", "lesson_table", "lesson_jnu_form"}:
        tpl = "lesson_simple"
    system_prompt = _build_lesson_system_prompt(req=req, full_context=full_context, kb_context=kb_context)

    last_flush = asyncio.get_event_loop().time()

    try:
        # meta
        meta_prompt = (
            'Return STRICT JSON only: {"targetAudience":"...","duration":"..."}.\n'
            'duration should be like "45分钟" (zh) or "45 min" (en).'
            if want_english
            else '仅输出严格 JSON：{"targetAudience":"...","duration":"..."}。\n'
                 'duration 形如 "45分钟"。'
        )
        meta = await _generate_lesson_section_with_llm(
            llm_settings=llm_settings,
            system_prompt=system_prompt,
            user_prompt=meta_prompt,
            temperature=0.3,
        )
        target_audience = str(meta.get("targetAudience") or ("Students" if want_english else "中学学生")).strip() or ("Students" if want_english else "中学学生")
        duration = str(meta.get("duration") or ("45 min" if want_english else "45分钟")).strip() or ("45 min" if want_english else "45分钟")

        # objectives
        if tpl == "lesson_jnu_form":
            obj_prompt = (
                'Return STRICT JSON only: {"objectives":["..."]}.\n'
                "Write 3-6 key points, then add ONE final item starting with 'Difficulty:'."
                if want_english
                else '仅输出严格 JSON：{"objectives":["..."]}。\n'
                     "先写 3~6 条“重点”，再额外补充 1 条以“难点：”开头的条目（作为最后一条）。"
            )
        else:
            obj_prompt = (
                'Return STRICT JSON only: {"objectives":["..."]}.\n'
                "Write 3-6 concise objectives."
                if want_english
                else '仅输出严格 JSON：{"objectives":["..."]}。\n'
                     "写 3~6 条可执行的教学目标。"
            )
        obj = await _generate_lesson_section_with_llm(
            llm_settings=llm_settings,
            system_prompt=system_prompt,
            user_prompt=obj_prompt,
            temperature=0.4,
        )
        objectives = obj.get("objectives")
        if not isinstance(objectives, list):
            objectives = []
        objectives = [str(x).strip() for x in objectives if str(x).strip()]
        if not objectives:
            objectives = _fallback_generate_lesson_plan(req).objectives
        elif tpl == "lesson_jnu_form":
            # 兜底修正：确保最后一条是“难点：/Difficulty:”
            diff = None
            rest: list[str] = []
            for item in objectives:
                text = str(item or "").strip()
                if not text:
                    continue
                is_diff = bool(re.match(r"^difficulty\s*[:：]", text, re.IGNORECASE)) if want_english else bool(re.match(r"^难点\s*[:：]", text))
                if is_diff and diff is None:
                    diff = text
                    continue
                rest.append(text)
            if diff is None:
                topic = (req.title or "").strip() or ("this lesson" if want_english else "本节课")
                diff = (
                    f"Difficulty: Explain and apply the key reasoning of {topic} in new contexts."
                    if want_english
                    else f"难点：掌握{topic}的关键推导与应用，并能迁移解决典型问题。"
                )
            objectives = rest + [diff]

        # 心跳：每10秒发一次注释，避免某些代理断连接
        now = asyncio.get_event_loop().time()
        if now - last_flush > 10:
            yield b": keep-alive\n\n"
            last_flush = now

        yield _encode_sse_data(json.dumps({"type": "section", "section": "objectives", "data": objectives}, ensure_ascii=False, separators=(",", ":")))

        # materials
        if tpl == "lesson_jnu_form":
            mat_prompt = (
                'Return STRICT JSON only: {"materials":["..."]}.\n'
                "List 4-10 items. Include at least one reference item (e.g., starts with 'Textbook:' or 'Reference:')."
                if want_english
                else '仅输出严格 JSON：{"materials":["..."]}。\n'
                     "列出 4~10 项教学材料/工具；并至少包含 1 条参考资料（例如以“教材：/参考书：/参考资料：”开头）。"
            )
        else:
            mat_prompt = (
                'Return STRICT JSON only: {"materials":["..."]}.\n'
                "List 3-8 materials/tools needed."
                if want_english
                else '仅输出严格 JSON：{"materials":["..."]}。\n'
                     "列出 3~8 项教学材料/工具。"
            )
        mat = await _generate_lesson_section_with_llm(
            llm_settings=llm_settings,
            system_prompt=system_prompt,
            user_prompt=mat_prompt,
            temperature=0.4,
        )
        materials = mat.get("materials")
        if not isinstance(materials, list):
            materials = []
        materials = [str(x).strip() for x in materials if str(x).strip()]
        if not materials:
            materials = _fallback_generate_lesson_plan(req).materials
        elif tpl == "lesson_jnu_form":
            ref_keywords = (
                ["textbook", "reference", "paper", "book", "isbn"]
                if want_english
                else ["教材", "课本", "参考", "文献", "论文", "书", "ISBN"]
            )
            has_ref = any(any(k.lower() in m.lower() for k in ref_keywords) for m in materials)
            if not has_ref:
                materials = materials + (["Textbook/Reference: ______"] if want_english else ["教材/参考书：______"])

        now = asyncio.get_event_loop().time()
        if now - last_flush > 10:
            yield b": keep-alive\n\n"
            last_flush = now

        yield _encode_sse_data(json.dumps({"type": "section", "section": "materials", "data": materials}, ensure_ascii=False, separators=(",", ":")))

        # procedure
        proc_prompt = (
            'Return STRICT JSON only: {"procedure":[{"step":"...","duration":"...","activity":"..."}]}.\n'
            "Write 4-8 steps; duration like '5 min'."
            if want_english
            else '仅输出严格 JSON：{"procedure":[{"step":"...","duration":"...","activity":"..."}]}。\n'
                 "写 4~8 步教学流程；duration 形如 '5分钟'。"
        )
        proc = await _generate_lesson_section_with_llm(
            llm_settings=llm_settings,
            system_prompt=system_prompt,
            user_prompt=proc_prompt,
            temperature=0.5,
        )
        procedure_raw = proc.get("procedure")
        procedure: list[LessonPlanProcedureStep] = []
        if isinstance(procedure_raw, list):
            for item in procedure_raw:
                if not isinstance(item, dict):
                    continue
                step = str(item.get("step") or "").strip()
                dur = str(item.get("duration") or "").strip()
                act = str(item.get("activity") or "").strip()
                if not step or not act:
                    continue
                if not dur:
                    dur = "5 min" if want_english else "5分钟"
                procedure.append(LessonPlanProcedureStep(step=step, duration=dur, activity=act))
        if not procedure:
            procedure = _fallback_generate_lesson_plan(req).procedure

        now = asyncio.get_event_loop().time()
        if now - last_flush > 10:
            yield b": keep-alive\n\n"
            last_flush = now

        yield _encode_sse_data(json.dumps({"type": "section", "section": "procedure", "data": [p.model_dump() for p in procedure]}, ensure_ascii=False, separators=(",", ":")))

        # homework
        if tpl == "lesson_jnu_form":
            hw_prompt = (
                'Return STRICT JSON only: {"homework":"..."}.\n'
                "Write questions/discussion/homework, concise but actionable."
                if want_english
                else '仅输出严格 JSON：{"homework":"..."}。\n'
                     "以“思考题/讨论题或作业”的形式输出，内容简洁可执行。"
            )
        else:
            hw_prompt = (
                'Return STRICT JSON only: {"homework":"..."}.\n'
                "Keep it concise."
                if want_english
                else '仅输出严格 JSON：{"homework":"..."}。\n'
                     "内容简洁可执行。"
            )
        hw = await _generate_lesson_section_with_llm(
            llm_settings=llm_settings,
            system_prompt=system_prompt,
            user_prompt=hw_prompt,
            temperature=0.4,
        )
        homework = str(hw.get("homework") or "").strip()
        if not homework:
            homework = _fallback_generate_lesson_plan(req).homework

        now = asyncio.get_event_loop().time()
        if now - last_flush > 10:
            yield b": keep-alive\n\n"
            last_flush = now

        yield _encode_sse_data(json.dumps({"type": "section", "section": "homework", "data": homework}, ensure_ascii=False, separators=(",", ":")))

        plan = LessonPlan(
            title=(req.title or "").strip() or ("Lesson Plan" if want_english else "教案"),
            targetAudience=target_audience,
            duration=duration,
            objectives=objectives,
            materials=materials,
            procedure=procedure,
            homework=homework,
        )
        yield _encode_sse_data(json.dumps({"type": "final", "data": plan.model_dump()}, ensure_ascii=False, separators=(",", ":")))
    except asyncio.CancelledError:
        logger.info("客户端已断开 /tools/lesson_plan SSE 连接，提前结束流")
        raise
    except Exception as exc:
        logger.error("LessonPlan 生成流异常: %s", exc, exc_info=True)
        payload = json.dumps({"type": "error", "text": f"教案生成异常：{exc}"}, ensure_ascii=False, separators=(",", ":"))
        yield _encode_sse_data(payload)
    finally:
        yield b"data: [DONE]\n\n"

async def iter_outline_text_chunks(
    prompt: str,
    language: str = "chinese",
    *,
    user_id: str = "default_user",
    metadata: dict[str, Any] | None = None,
):
    """
    抽象出大纲 Agent 的文本增量迭代器：
    - 只关心 chunk_data["type"] == "text" 的部分
    - 统一日志与空文本过滤
    """
    outline_wrapper = A2AOutlineClientWrapper(session_id=uuid.uuid4().hex, agent_url=_get_outline_api())
    async for chunk_data in outline_wrapper.generate(prompt, language=language, user_id=user_id, metadata=metadata):
        logger.info(f"生成大纲输出的chunk_data: {chunk_data}")
        chunk_type = chunk_data.get("type")

        if chunk_type == "text":
            text = chunk_data.get("text") or ""
            if not text:
                continue
            yield text


async def stream_outline_sse(
    prompt: str,
    language: str = "chinese",
    *,
    user_id: str = "default_user",
    metadata: dict[str, Any] | None = None,
):
    """
    将大纲 Agent 的响应以 SSE 形式向前端流式输出。
    - media_type: text/event-stream
    - 每个 chunk_data["text"] 作为一条 SSE 事件发送
    - 如 text 内部包含换行，按 SSE 规范拆成多行 data:
    - 结束时发送 data: [DONE]
    """
    async for text in iter_outline_text_chunks(prompt, language, user_id=user_id, metadata=metadata):
            # 按行拆分，遵守 SSE 规范：一条事件内多行 data:
            lines = text.splitlines()
            # 如果 text 以换行结尾，splitlines() 会吞掉末尾空行，这里记录一下
            has_trailing_newline = text.endswith("\n")

            if not lines:
                # 纯换行的情况，例如 text == "\n" 或 "\n\n"
                # 用一个空 data 行表示，然后前端按照事件边界追加换行
                yield b"data:\n\n"
            else:
                for line in lines:
                    # 每一行作为一条 data: 行
                    yield f"data: {line}\n".encode("utf-8")
                if has_trailing_newline:
                    # 保留结尾换行：再补一个空 data 行
                    yield b"data:\n"
            # 事件结束
            yield b"\n"

    # 目前前端不依赖 artifact/metadata/final 这几类事件，这里仅记录日志即可
    # 如果后续需要，可以在此扩展不同类型的 SSE 事件

    # 显式结束信号，前端据此收尾
    yield b"data: [DONE]\n\n"


def _encode_sse_data(payload: str) -> bytes:
    """
    将任意文本安全编码为 SSE data 事件，兼容多行内容。
    """
    if payload is None:
        payload = ""

    lines = payload.splitlines()
    has_trailing_newline = payload.endswith("\n")

    if not lines:
        return b"data:\n\n"

    chunks: list[bytes] = []
    for line in lines:
        chunks.append(f"data: {line}\n".encode("utf-8"))
    if has_trailing_newline:
        chunks.append(b"data:\n")
    chunks.append(b"\n")
    return b"".join(chunks)


@app.post("/tools/aippt_outline")
async def aippt_outline(request: AipptRequest):
    assert request.stream, "只支持流式的返回大纲"
    logger.info(f"前端*outline***=====>用户输入：{request.language}")
    async def event_generator():
        async for chunk in stream_outline_sse(request.content, request.language):
            yield chunk

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.post("/tools/outline")
@app.post("/tools/aippt_outline_unified")
async def aippt_outline_unified(
    content: str = Form(None),           # 主题文本（可选）
    file: UploadFile = File(None),       # 上传文件（可选）
    language: str = Form("chinese"),
    user_id: str = Form("default_user"),
    folder_id: int | str = Form(0),
    file_type: str | None = Form(None),
    kb_file_ids: list[str] | None = Form(None),  # 可选：限定从哪些 KB 文件检索参考片段
    outline_length: str = Form("standard"),  # short | standard | long
    use_web_search: bool = Form(True),
):
    """
    统一的大纲生成 API，支持两种模式：
    - 主题模式：只传 content，根据主题生成大纲
    - 文档模式：传 file，解析文档后生成大纲
    - 混合模式：同时传 content 和 file，以文档为主，主题作为补充上下文
    """
    content_text = (content or "").strip()
    has_content = bool(content_text)
    has_file = file is not None

    if not has_content and not has_file:
        raise HTTPException(status_code=400, detail="请提供主题或文件")

    file_content = ""
    personaldb_url = _get_personaldb_url()

    # 如果有文件，先解析文件内容
    if has_file:
        if not personaldb_url:
            raise HTTPException(status_code=500, detail="PERSONAL_DB 未配置")

        # 生成 fileId
        file_id = str(int(time.time() * 1000))

        # 推断 fileType
        actual_file_type = file_type
        if not actual_file_type and file.filename and "." in file.filename:
            actual_file_type = file.filename.rsplit(".", 1)[-1]

        # 组装请求数据
        data = {
            "userId": str(user_id),
            "fileId": file_id,
            "folderId": str(folder_id),
        }
        if actual_file_type:
            data["fileType"] = actual_file_type

        # 读取文件内容
        file_bytes = await file.read()
        if not file_bytes:
            raise HTTPException(status_code=400, detail="文件内容为空")

        files_payload = {
            "file": (
                file.filename or "uploaded_file",
                file_bytes,
                file.content_type or "application/octet-stream",
            )
        }

        upload_url = f"{personaldb_url}/upload/"

        # 内部服务调用（personaldb）不应受系统代理环境变量影响
        async with httpx.AsyncClient(trust_env=False, timeout=httpx.Timeout(360.0)) as client:
            try:
                resp = await client.post(
                    upload_url,
                    data=data,
                    files=files_payload,
                )
                if resp.status_code >= 400:
                    logger.info(f"[personaldb {resp.status_code}] {resp.text}")
                    resp.raise_for_status()

                try:
                    result = resp.json()
                except ValueError:
                    raise HTTPException(status_code=502, detail=f"personaldb 返回的不是 JSON：{resp.text}")

                markdown_content = result.get("markdown_content")
                if markdown_content is None:
                    raise HTTPException(status_code=500, detail="personaldb 响应缺少 'markdown_content'")

                file_content = markdown_content

            except httpx.TimeoutException:
                raise HTTPException(status_code=504, detail="Request to personaldb timed out.")
            except httpx.HTTPStatusError as exc:
                raise HTTPException(status_code=exc.response.status_code, detail=exc.response.text)
            except httpx.RequestError as exc:
                raise HTTPException(status_code=500, detail=f"Error connecting to personaldb: {exc}")

    full_context = ""
    kb_context = ""
    resolved_kb_file_ids = _normalize_kb_file_ids(kb_file_ids)
    if resolved_kb_file_ids:
        if personaldb_url and await _is_personaldb_ready(personaldb_url):
            full_context, kb_context = await _build_personaldb_kb_contexts(
                personaldb_url,
                user_id=str(user_id),
                query=content_text,
                kb_file_ids=resolved_kb_file_ids,
                rag_topk=5,
            )
        else:
            logger.info("personaldb 不可用，跳过 kb_file_ids 检索增强：%s", personaldb_url)

    prompt_parts: list[str] = []
    if content_text:
        prompt_parts.append(content_text)
    if file_content:
        prompt_parts.append(f"参考文档内容（来自你上传的文件）：\n{file_content}")
    if full_context:
        prompt_parts.append(
            "课程产出（全文，不经检索）：\n"
            "（说明：仅供参考用于一致性对齐，不要原文照抄，也不要把其目录/结构当作本次大纲结构。）\n"
            f"{full_context}"
        )
    if kb_context:
        prompt_parts.append(f"参考资料检索片段（RAG）：\n{kb_context}")
    prompt = "\n\n".join(prompt_parts)

    outline_length_norm = (outline_length or "").strip().lower() or "standard"
    if outline_length_norm not in {"short", "standard", "long"}:
        logger.info("outline_length 非法，回落为 standard: %s", outline_length)
        outline_length_norm = "standard"

    outline_metadata = {
        "outline_length": outline_length_norm,
        "use_web_search": bool(use_web_search),
        "user_id": str(user_id),
    }

    logger.info(
        "统一大纲API*outline***=====>：language=%s, has_file=%s, has_content=%s, outline_length=%s, use_web_search=%s",
        language,
        has_file,
        has_content,
        outline_length_norm,
        bool(use_web_search),
    )

    async def event_generator():
        async for chunk in stream_outline_sse(
            prompt,
            language,
            user_id=str(user_id),
            metadata=outline_metadata,
        ):
            yield chunk

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.post("/tools/outline_from_file")
@app.post("/tools/aippt_outline_from_file")
async def aippt_outline_from_file(
    user_id: int|str = Form(...),
    file: UploadFile = File(None),  # 允许缺省，这样我们可以决定走 file 或 url
    url: str | None = Form(None),
    folder_id: int|str = Form(0),
    file_type: str | None = Form(None),
    language: str = Form("chinese"),  # 添加language参数，默认为chinese
):
    """
    对齐 personaldb 的 /upload/：
    - 必填: userId, fileId
    - 可选: folderId (默认0), fileType
    - file 与 url 互斥，至少一个
    """
    personaldb_api_url = os.getenv("PERSONAL_DB")
    if not personaldb_api_url:
        raise HTTPException(status_code=500, detail="PERSONAL_DB 未配置")

    # 互斥校验（与 personaldb 完全一致）
    has_file = file is not None
    has_url = bool(url and url.strip())

    # 生成 fileId（字符串更稳；personaldb 会 int()）
    file_id = str(int(time.time() * 1000))

    # 推断 fileType（当上传文件时且未显式传入）
    if has_file and not file_type:
        if file.filename and "." in file.filename:
            file_type = file.filename.rsplit(".", 1)[-1]
        else:
            file_type = "unknown"

    # 组装 multipart/form-data
    # 注意：即使是 url 分支，也仍用 multipart，personaldb 也能解析 form
    data = {
        "userId": str(user_id),
        "fileId": file_id,
        "folderId": str(folder_id),
    }
    if file_type:
        data["fileType"] = file_type
    if has_url:
        data["url"] = url.strip()

    files_payload = None
    if has_file:
        # 读取一次到内存，httpx 需要 (filename, bytes/obj, content_type)
        file_bytes = await file.read()
        if not file_bytes:
            raise HTTPException(status_code=400, detail="文件内容为空")
        files_payload = {
            "file": (
                file.filename or "uploaded_file",
                file_bytes,
                file.content_type or "application/octet-stream",
            )
        }

    upload_url = f"{personaldb_api_url.rstrip('/')}/upload/"

    # 内部服务调用（personaldb）不应受系统代理环境变量影响
    async with httpx.AsyncClient(trust_env=False) as client:
        try:
            resp = await client.post(
                upload_url,
                data=data,
                files=files_payload,
                timeout=360.0,
            )
            # 不直接 raise，先打日志方便定位
            if resp.status_code >= 400:
                # 打印下游返回体，personaldb 对错误信息写得很清楚
                logger.info(f"[personaldb {resp.status_code}] {resp.text}")
                resp.raise_for_status()

            # personaldb 的处理函数最终会返回一个 JSON（你上游期望里要有 markdown_content）
            try:
                result = resp.json()
            except ValueError:
                raise HTTPException(status_code=502, detail=f"personaldb 返回的不是 JSON：{resp.text}")

            markdown_content = result.get("markdown_content")
            if markdown_content is None:
                raise HTTPException(status_code=500, detail="personaldb 响应缺少 'markdown_content'")
            logger.info(f"本地上传文件*outline***=====>：{ {'language': language} }")

            async def event_generator():
                async for chunk in stream_outline_sse(markdown_content, language):
                    yield chunk

            return StreamingResponse(
                event_generator(),
                media_type="text/event-stream",
                headers={
                    "Cache-Control": "no-cache, no-transform",
                    "Connection": "keep-alive",
                    "X-Accel-Buffering": "no",
                },
            )

        except httpx.TimeoutException:
            raise HTTPException(status_code=504, detail="Request to personaldb timed out.")
        except httpx.HTTPStatusError as exc:
            # 透传 personaldb 的错误详情，便于你在日志里看到具体字段问题
            raise HTTPException(status_code=exc.response.status_code, detail=exc.response.text)
        except httpx.RequestError as exc:
            raise HTTPException(status_code=500, detail=f"Error connecting to personaldb: {exc}")

class AipptContentRequest(BaseModel):
    content: str
    language: str = "zh"  #默认中文
    sessionId: str = ""  # 当使用知识库时，需要根据用户的user_id查询对应的知识库
    generateFromUploadedFile: bool = False  # 是否从上传的文件中生成PPT内容
    generateFromWebSearch: bool = True  # 是否从网络搜索中生成PPT内容
    generateWithImages: bool = False  # 是否启用“联网配图”（开启：检索网络图片；关闭：使用预设图片池）
    kb_folder_ids: list[int] | None = None  # 仅当启用知识库检索时生效，用于过滤可检索的 folder_id
    kb_file_ids: list[str] | None = None  # 仅当启用知识库检索时生效，用于过滤可检索的 file_id（更精确）

def _kb_ok(data):
    return {"ok": True, "data": data}


def _kb_error(code: str, message: str, status_code: int = 500):
    return JSONResponse(
        status_code=status_code,
        content={"ok": False, "error": {"code": code, "message": message}},
    )


def _kb_safe_filename(name: str) -> str:
    # 避免 header 注入与路径穿越
    safe = (name or "").strip().replace("\\", "_").replace("/", "_")
    safe = safe.replace("\r", "_").replace("\n", "_").replace("\t", "_")
    return safe


def _kb_build_export_filename(file_name: str, file_type: str, file_id: str) -> str:
    base = _kb_safe_filename(file_name) or _kb_safe_filename(file_id) or "export"
    base = base or "export"

    lower = base.lower()
    if "." in Path(base).name:
        # 已带后缀：若不是可读文本后缀，则追加 .md
        if lower.endswith((".md", ".txt")):
            return base
        return f"{base}.md"

    ext = (file_type or "").strip().lower().lstrip(".")
    if ext in {"md", "markdown"}:
        return f"{base}.md"
    if ext in {"txt", "text"}:
        return f"{base}.txt"
    if ext and ext not in {"unknown"}:
        return f"{base}.{ext}.md"
    return f"{base}.md"


def _artifact_safe_filename(name: str) -> str:
    """
    artifacts 落盘用的文件名净化：
    - 兼容 Windows/WSL 的 NTFS（例如 ':' 在文件名中不合法）
    - 避免路径穿越
    """
    safe = _kb_safe_filename(name)
    for ch in [":", "<", ">", '"', "|", "?", "*"]:
        safe = safe.replace(ch, "_")
    safe = safe.strip()
    if safe in {"", ".", ".."}:
        return ""
    return safe


def _artifact_safe_segment(value: str) -> str:
    safe = _artifact_safe_filename(str(value or ""))
    safe = safe.replace("..", "_").strip("._")
    return safe or "unknown"


def _get_artifact_root_dir() -> Path:
    """
    Artifacts 根目录：
    - env: TEACHDO_ARTIFACT_DIR（默认 var/artifacts）
    - 相对路径按 repo root 解析（复用 _find_repo_root）
    """
    configured = (os.environ.get("TEACHDO_ARTIFACT_DIR") or "").strip()
    repo_root = _find_repo_root(Path(__file__).resolve())
    root = (Path(configured) if configured else Path("var/artifacts"))
    if not root.is_absolute():
        root = repo_root / root
    root.mkdir(parents=True, exist_ok=True)
    return root


def _normalize_artifact_kind(kind: str) -> str | None:
    k = (kind or "").strip().lower()
    return k if k in {"pptx", "docx"} else None


def _artifact_media_type(kind: str) -> str:
    k = (kind or "").strip().lower()
    if k == "pptx":
        return "application/vnd.openxmlformats-officedocument.presentationml.presentation"
    if k == "docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return "application/octet-stream"


def _get_artifact_material_dir(*, user_id: str, material_id: str) -> Path:
    root = _get_artifact_root_dir()
    u = _artifact_safe_segment(user_id)
    m = _artifact_safe_segment(material_id)
    path = root / u / m
    path.mkdir(parents=True, exist_ok=True)
    return path


def _artifact_index_path(material_dir: Path) -> Path:
    return material_dir / "index.json"


def _load_artifact_index(material_dir: Path) -> list[dict[str, Any]]:
    path = _artifact_index_path(material_dir)
    if not path.exists():
        return []
    try:
        raw = path.read_text("utf-8")
        obj = json.loads(raw)
    except Exception:
        logger.warning("读取 artifact 索引失败: %s", path, exc_info=True)
        return []

    items: Any = None
    if isinstance(obj, dict):
        items = obj.get("artifacts")
    elif isinstance(obj, list):
        items = obj
    if not isinstance(items, list):
        return []
    return [it for it in items if isinstance(it, dict)]


def _write_artifact_index(material_dir: Path, items: list[dict[str, Any]]) -> None:
    path = _artifact_index_path(material_dir)
    tmp = path.with_suffix(".json.tmp")
    payload = {"artifacts": items}
    tmp.write_text(json.dumps(payload, ensure_ascii=False, separators=(",", ":")), "utf-8")
    os.replace(tmp, path)


def _artifact_public_meta(item: dict[str, Any]) -> dict[str, Any]:
    meta: dict[str, Any] = {
        "artifact_id": str(item.get("artifact_id") or ""),
        "kind": str(item.get("kind") or ""),
        "file_name": str(item.get("file_name") or ""),
    }
    created_at = item.get("created_at")
    if created_at is not None:
        try:
            meta["created_at"] = int(created_at)
        except Exception:
            logger.debug("忽略非法 created_at: %r", created_at, exc_info=True)
            pass
    size = item.get("size")
    if size is not None:
        try:
            meta["size"] = int(size)
        except Exception:
            logger.debug("忽略非法 size: %r", size, exc_info=True)
            pass
    return meta


def _save_artifact_bytes(
    *,
    user_id: str,
    material_id: str,
    kind: str,
    file_bytes: bytes,
    file_name: str,
) -> dict[str, Any]:
    normalized_kind = _normalize_artifact_kind(kind)
    if not normalized_kind:
        raise ValueError(f"非法 kind：{kind}")

    material_dir = _get_artifact_material_dir(user_id=str(user_id), material_id=str(material_id))

    artifact_id = uuid.uuid4().hex
    safe_name = _artifact_safe_filename(file_name) or f"{normalized_kind}.{normalized_kind}"
    if not safe_name.lower().endswith(f".{normalized_kind}"):
        safe_name = safe_name + f".{normalized_kind}"

    stored_name = f"{artifact_id}__{safe_name}"
    stored_path = material_dir / stored_name
    stored_path.write_bytes(file_bytes)

    created_at = int(time.time() * 1000)
    size = len(file_bytes) if isinstance(file_bytes, (bytes, bytearray)) else 0

    items = _load_artifact_index(material_dir)
    items.append(
        {
            "artifact_id": artifact_id,
            "kind": normalized_kind,
            "file_name": safe_name,
            "stored_name": stored_name,
            "created_at": created_at,
            "size": size,
        }
    )
    _write_artifact_index(material_dir, items)
    return _artifact_public_meta(items[-1])


def _get_personaldb_url() -> str | None:
    url = os.environ.get("PERSONAL_DB")
    return url.rstrip("/") if url else None


async def _is_personaldb_ready(personaldb_url: str) -> bool:
    try:
        async with httpx.AsyncClient(trust_env=False, timeout=httpx.Timeout(2.0)) as client:
            resp = await client.get(f"{personaldb_url}/healthz")
            return resp.status_code == 200
    except Exception:
        logger.debug("检查 personaldb 健康状态失败", exc_info=True)
        return False


def _normalize_kb_file_ids(kb_file_ids: list[str] | None) -> list[str]:
    """
    归一化 kb_file_ids：
    - 去空白
    - 去重（保持稳定顺序）
    """
    if not isinstance(kb_file_ids, list) or not kb_file_ids:
        return []
    seen: set[str] = set()
    resolved: list[str] = []
    for raw in kb_file_ids:
        sid = str(raw).strip()
        if not sid or sid in seen:
            continue
        seen.add(sid)
        resolved.append(sid)
    return resolved


def _split_kb_file_ids(kb_file_ids: list[str]) -> tuple[list[str], list[str]]:
    """
    按约定将 KB 文件拆分为两类：
    - full_ids：全文注入（gen:/full: 前缀）→ 拉取全文加入上下文（不经检索）
    - rag_ids：RAG 检索（非 gen:/full:）→ 仅用于 personaldb /search（只注入相关片段）
    """
    full_ids: list[str] = []
    rag_ids: list[str] = []
    for fid in kb_file_ids or []:
        if str(fid).startswith("gen:") or str(fid).startswith("full:"):
            full_ids.append(str(fid))
        else:
            rag_ids.append(str(fid))
    return full_ids, rag_ids


def _format_personaldb_search_context(
    result: object,
    *,
    max_chunks: int = 5,
    max_total_chars: int = 4000,
    max_chunk_chars: int = 800,
) -> str:
    """
    将 personaldb /search 的返回格式化为可拼进 prompt 的参考内容。
    约束：
    - 限制总长度，避免 prompt 过大导致模型效果变差或超限
    - 每个 chunk 截断
    """
    if not isinstance(result, dict):
        return ""

    documents = result.get("documents")
    metadatas = result.get("metadatas")
    distances = result.get("distances")

    docs_row: list[object] = []
    metas_row: list[object] = []
    dists_row: list[object] = []

    if isinstance(documents, list) and documents and isinstance(documents[0], list):
        docs_row = documents[0]
    if isinstance(metadatas, list) and metadatas and isinstance(metadatas[0], list):
        metas_row = metadatas[0]
    if isinstance(distances, list) and distances and isinstance(distances[0], list):
        dists_row = distances[0]

    blocks: list[str] = []
    total = 0
    for idx, doc in enumerate(docs_row):
        if not isinstance(doc, str):
            continue
        text = doc.strip()
        if not text:
            continue

        meta = metas_row[idx] if idx < len(metas_row) else None
        file_id = ""
        file_name = ""
        folder_id = None
        if isinstance(meta, dict):
            file_id = str(meta.get("file_id") or meta.get("fileId") or "").strip()
            file_name = str(meta.get("file_name") or meta.get("fileName") or "").strip()
            folder_id = meta.get("folder_id") if meta.get("folder_id") is not None else meta.get("folderId")

        dist = dists_row[idx] if idx < len(dists_row) else None
        dist_str = ""
        if dist is not None:
            try:
                dist_str = f"{float(dist):.4f}"
            except Exception:
                dist_str = ""

        if len(text) > max_chunk_chars:
            text = text[:max_chunk_chars].rstrip() + "…"

        meta_bits: list[str] = []
        if file_name:
            meta_bits.append(file_name)
        if file_id:
            meta_bits.append(f"file_id={file_id}")
        if folder_id is not None:
            try:
                meta_bits.append(f"folder_id={int(folder_id)}")
            except Exception:
                logger.debug("忽略非法 folder_id: %r", folder_id, exc_info=True)
                pass
        if dist_str:
            meta_bits.append(f"distance={dist_str}")
        meta_line = " / ".join(meta_bits).strip() or "KB chunk"

        block = f"[{len(blocks) + 1}] {meta_line}\n{text}"
        if total + len(block) > max_total_chars:
            break
        blocks.append(block)
        total += len(block) + 2
        if len(blocks) >= max_chunks:
            break

    return "\n\n".join(blocks).strip()


async def _search_personaldb_kb_context(
    personaldb_url: str,
    *,
    user_id: str,
    query: str,
    kb_file_ids: list[str],
    topk: int = 5,
) -> str:
    """
    从 personaldb 检索知识库片段，作为大纲生成的参考上下文。
    注意：检索失败时不应阻断主流程（仍可用主题生成大纲）。
    """
    if not query.strip():
        return ""
    if not kb_file_ids:
        return ""

    payload = {
        "userId": str(user_id),
        "query": str(query),
        "keyword": "",
        "topk": int(topk),
        "fileIds": list(kb_file_ids),
    }

    url = f"{personaldb_url}/search"
    async with httpx.AsyncClient(trust_env=False, timeout=httpx.Timeout(20.0)) as client:
        try:
            resp = await client.post(url, json=payload)
            if resp.status_code >= 400:
                logger.info("[personaldb %s] %s", resp.status_code, resp.text)
                return ""
            try:
                result = resp.json()
            except ValueError:
                logger.info("personaldb /search 返回非 JSON：%s", resp.text)
                return ""
        except Exception as exc:
            logger.info("personaldb /search 调用失败：%s", exc)
            return ""

    return _format_personaldb_search_context(result)


async def _load_personaldb_full_text_context(
    personaldb_url: str,
    *,
    user_id: str,
    file_ids: list[str],
    max_file_chars: int = 40_000,
    max_total_chars: int = 120_000,
) -> str:
    """
    从 personaldb 拉取指定 file_ids 的全文内容（不经检索），并拼接为可注入 prompt 的上下文。

    约束（防止 prompt 过大）：
    - 单文件最多 max_file_chars 字符
    - 总计最多 max_total_chars 字符
    - 超限截断时在上下文中标注“已截断”
    """
    if not file_ids:
        return ""

    def _build_prefix(index: int, title: str, meta_line: str, notes: list[str]) -> str:
        lines = [f"[{index}] {title}"]
        if meta_line:
            lines.append(meta_line)
        if notes:
            lines.append("（" + "；".join([x for x in notes if x]) + "）")
        return "\n".join(lines).rstrip() + "\n"

    blocks: list[str] = []
    total_chars = 0  # 仅统计返回字符串长度（含分隔空行）

    async with httpx.AsyncClient(trust_env=False, timeout=httpx.Timeout(20.0)) as client:
        for fid in file_ids:
            if total_chars >= max_total_chars:
                break

            url = f"{personaldb_url}/files/{user_id}/{fid}/content"
            try:
                resp = await client.get(url)
            except Exception as exc:
                logger.info("personaldb /files/.../content 调用失败：%s", exc)
                continue

            if resp.status_code == 404:
                continue
            if resp.status_code >= 400:
                logger.info("[personaldb %s] %s", resp.status_code, resp.text)
                continue

            try:
                payload = resp.json()
            except ValueError:
                logger.info("personaldb /files/.../content 返回非 JSON：%s", resp.text)
                continue

            if not isinstance(payload, dict):
                continue

            content = payload.get("content")
            if not isinstance(content, str):
                continue
            text = content.strip()
            if not text:
                continue

            file_name = str(payload.get("file_name") or payload.get("fileName") or "").strip()
            file_type = str(payload.get("file_type") or payload.get("fileType") or "").strip()

            title = file_name or fid
            notes: list[str] = []

            if len(text) > max_file_chars:
                text = text[:max_file_chars].rstrip() + "…"
                notes.append(f"已按单文件上限截断（{max_file_chars} chars）")

            meta_bits = [f"file_id={fid}"]
            if file_type:
                meta_bits.append(f"type={file_type}")
            meta_line = " / ".join(meta_bits).strip()

            index = len(blocks) + 1

            # 预留分隔空行：除首块外，每块前面会有 "\n\n"
            remaining_total = max_total_chars - total_chars
            if blocks:
                remaining_total -= 2
            if remaining_total <= 0:
                break

            prefix = _build_prefix(index, title, meta_line, notes)
            remaining_for_content = remaining_total - len(prefix)
            if remaining_for_content <= 0:
                break

            if len(text) > remaining_for_content:
                if "已按总长度上限截断" not in notes:
                    notes.append("已按总长度上限截断")
                prefix = _build_prefix(index, title, meta_line, notes)
                remaining_for_content = remaining_total - len(prefix)
                if remaining_for_content <= 0:
                    break
                if len(text) > remaining_for_content:
                    # 至少留 1 个字符给省略号
                    cut = max(0, remaining_for_content - 1)
                    text = text[:cut].rstrip() + "…"

            block = (prefix + text).strip()
            if blocks:
                total_chars += 2
            blocks.append(block)
            total_chars += len(block)

    return "\n\n".join(blocks).strip()


async def _build_personaldb_kb_contexts(
    personaldb_url: str,
    *,
    user_id: str,
    query: str,
    kb_file_ids: list[str],
    rag_topk: int = 5,
) -> tuple[str, str]:
    """
    给定“用户当前选中的 kb_file_ids”，按约定构建两类上下文：
    - full_context：gen: 全文注入（不经检索）
    - rag_context：非 gen: 走 /search 的检索片段
    """
    full_ids, rag_ids = _split_kb_file_ids(kb_file_ids)
    full_context = ""
    rag_context = ""
    if rag_ids:
        rag_context = await _search_personaldb_kb_context(
            personaldb_url,
            user_id=str(user_id),
            query=str(query),
            kb_file_ids=rag_ids,
            topk=int(rag_topk),
        )
    if full_ids:
        full_context = await _load_personaldb_full_text_context(
            personaldb_url,
            user_id=str(user_id),
            file_ids=full_ids,
        )
    return full_context, rag_context


def _pick_last_user_message(messages: list[AssistantChatMessage]) -> str:
    """
    从历史消息中找到最后一条 user 消息，作为 RAG 检索 query。
    """
    for msg in reversed(messages or []):
        if msg.role == "user" and (msg.content or "").strip():
            return msg.content.strip()
    return ""


def _build_assistant_system_prompt(
    *,
    material: AssistantMaterialContext | None,
    full_context: str,
    kb_context: str,
    language: str,
) -> str:
    """
    构建助教的 system prompt：
    - 教学资料上下文（标题/学科/简介/目标）
    - KB 检索片段（若有）
    """
    lang = (language or "zh").strip().lower()
    want_english = lang in {"en", "english"}

    if want_english:
        base = (
            "You are TeachDo's AI teaching assistant.\n"
            "Your goal is to help teachers design lessons, explain concepts, generate exercises, and answer questions.\n"
            "Rules:\n"
            "- Be accurate and practical.\n"
            "- If the user question is ambiguous, ask 1-2 clarifying questions.\n"
            "- When course outputs (full text) are provided, prioritize them as context.\n"
            "- When KB snippets (RAG) are provided, use them as grounding; if insufficient, say so.\n"
            "- Use concise bullets/steps when helpful.\n"
        )
    else:
        base = (
            "你是 TeachDo 的 AI 教学助教。\n"
            "你的目标是帮助教师进行教学设计、知识点讲解、题目生成与答疑。\n"
            "规则：\n"
            "- 回答要准确、可操作。\n"
            "- 问题不清晰时，先问 1~2 个澄清问题。\n"
            "- 若提供了课程产出全文（不经检索），应优先基于全文作答。\n"
            "- 若提供了参考资料检索片段（RAG），应优先基于片段作答；片段不足时要明确说明。\n"
            "- 需要时用条目/步骤输出。\n"
        )

    context_bits: list[str] = []
    if material and (material.title or "").strip():
        title = material.title.strip()
        if want_english:
            context_bits.append(f"Current teaching material: {title}")
        else:
            context_bits.append(f"当前教学资料：{title}")
        if material.subject:
            context_bits.append(("Subject: " if want_english else "学科：") + str(material.subject).strip())
        if material.description:
            context_bits.append(("Description: " if want_english else "简介：") + str(material.description).strip())
        if material.objectives:
            context_bits.append(("Objectives: " if want_english else "教学目标：") + str(material.objectives).strip())

    if full_context and full_context.strip():
        if want_english:
            context_bits.append("Course outputs (full text, not retrieved):\n" + full_context.strip())
        else:
            context_bits.append("课程产出（全文，不经检索）：\n" + full_context.strip())

    if kb_context and kb_context.strip():
        if want_english:
            context_bits.append("Reference snippets (RAG):\n" + kb_context.strip())
        else:
            context_bits.append("参考资料检索片段（RAG）：\n" + kb_context.strip())

    if not context_bits:
        return base

    return base + "\n\n" + "\n".join(context_bits).strip()


def _get_assistant_llm_settings() -> dict[str, str]:
    """
    助教模型配置：
    - 优先读取 ASSISTANT_*，若未配置则回退到 OUTLINE_*（保持与计划一致，减少新增配置成本）。
    - 当前仅支持 openai 协议（OpenAI 兼容 base_url）。
    """
    llm_type = (os.getenv("ASSISTANT_TYPE") or os.getenv("OUTLINE_TYPE") or "").strip().lower()
    llm_model = (os.getenv("ASSISTANT_MODEL") or os.getenv("OUTLINE_MODEL") or "").strip()
    llm_api_key = (os.getenv("ASSISTANT_API_KEY") or os.getenv("OUTLINE_API_KEY") or "").strip()
    llm_base_url = (os.getenv("ASSISTANT_BASE_URL") or os.getenv("OUTLINE_BASE_URL") or "https://api.openai.com/v1").strip()

    if not llm_type or not llm_model:
        raise RuntimeError("助教模型未配置，请设置 ASSISTANT_TYPE/ASSISTANT_MODEL（或复用 OUTLINE_* 配置）")
    openai_compatible = {"openai", "ollama", "vllm", "local_openai", "xinference"}
    if llm_type not in openai_compatible:
        raise RuntimeError(f"当前助教仅支持 openai 兼容协议，检测到 ASSISTANT_TYPE/OUTLINE_TYPE={llm_type}")
    if not llm_api_key:
        raise RuntimeError("缺少助教模型 API Key，请设置 ASSISTANT_API_KEY（或复用 OUTLINE_API_KEY）")

    return {"type": llm_type, "model": llm_model, "api_key": llm_api_key, "base_url": llm_base_url}


async def iter_assistant_text_chunks(
    *,
    model: str,
    api_key: str,
    base_url: str,
    messages: list[dict[str, str]],
    temperature: float = 0.6,
) -> AsyncGenerator[str, None]:
    """
    通过 OpenAI 兼容协议调用 Chat Completions，并将增量 token 作为文本片段 yield。
    注意：这里不直接向客户端暴露上游 SSE，避免不同兼容网关的格式差异影响前端解析。
    """
    url = f"{base_url.rstrip('/')}/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": model,
        "messages": messages,
        "temperature": temperature,
        "stream": True,
    }

    timeout = httpx.Timeout(connect=60.0, write=60.0, pool=60.0, read=None)
    async with httpx.AsyncClient(timeout=timeout, trust_env=False) as client:
        async with client.stream("POST", url, headers=headers, json=payload) as resp:
            if resp.status_code >= 400:
                text = await resp.aread()
                raise RuntimeError(f"LLM 请求失败：{resp.status_code} {text.decode('utf-8', errors='ignore')}")

            async for line in resp.aiter_lines():
                if not line:
                    continue
                if line.startswith(":"):
                    continue
                if not line.startswith("data:"):
                    continue
                data = line[len("data:") :].strip()
                if not data:
                    continue
                if data == "[DONE]":
                    break
                try:
                    obj = json.loads(data)
                except Exception:
                    # 某些网关可能夹杂非 JSON 行，忽略
                    logger.debug("忽略非 JSON SSE 数据行: %r", data, exc_info=True)
                    continue

                # OpenAI 标准：choices[0].delta.content
                delta = None
                try:
                    choices = obj.get("choices") if isinstance(obj, dict) else None
                    if isinstance(choices, list) and choices:
                        choice0 = choices[0] if isinstance(choices[0], dict) else {}
                        delta_obj = choice0.get("delta") if isinstance(choice0.get("delta"), dict) else {}
                        delta = delta_obj.get("content")
                        if not delta:
                            # 兼容少数实现：message.content 直接返回完整段落
                            msg_obj = choice0.get("message") if isinstance(choice0.get("message"), dict) else {}
                            delta = msg_obj.get("content")
                except Exception:
                    logger.debug("提取 SSE delta 失败", exc_info=True)
                    delta = None

                if isinstance(delta, str) and delta:
                    yield delta


async def stream_assistant_sse(
    *,
    llm_settings: dict[str, str],
    messages: list[dict[str, str]],
) -> AsyncGenerator[bytes, None]:
    """
    将助教回答以 SSE 向前端流式输出（data: 增量文本，结束 data: [DONE]）。
    若发生异常，发送一条 type=error 的 JSON 事件后结束。
    """
    try:
        async for delta in iter_assistant_text_chunks(
            model=llm_settings["model"],
            api_key=llm_settings["api_key"],
            base_url=llm_settings["base_url"],
            messages=messages,
        ):
            yield _encode_sse_data(delta)
    except asyncio.CancelledError:
        logger.info("客户端已断开 /tools/assistant_chat SSE 连接，提前结束流")
        raise
    except Exception as exc:
        logger.error("助教对话流异常: %s", exc, exc_info=True)
        payload = json.dumps(
            {
                "type": "error",
                "text": f"助教服务异常：{exc}",
            },
            ensure_ascii=False,
        )
        yield _encode_sse_data(payload)
    finally:
        yield b"data: [DONE]\n\n"


async def stream_content_response(
    markdown_content: str,
    language,
    generateFromUploadedFile,
    generateFromWebSearch,
    user_id,
    generateWithImages: bool = False,
    kb_folder_ids: list[int] | None = None,
    kb_file_ids: list[str] | None = None,
):
    match = re.search(r"(# .*)", markdown_content, flags=re.DOTALL)
    result = markdown_content[match.start():] if match else markdown_content
    logger.info(f"用户输入的markdown大纲是：{result}")

    content_wrapper = A2AContentClientWrapper(session_id=uuid.uuid4().hex, agent_url=_get_content_api())

    search_engine = []
    if generateFromUploadedFile:
        search_engine.append("KnowledgeBaseSearch")
    if generateFromWebSearch:
        search_engine.append("DocumentSearch")
    # 联网配图：由 Writer 直接调用 SearchImage 决定 query（失败时服务端再做兜底注入）
    if bool(generateWithImages) and "SearchImage" not in search_engine:
        search_engine.append("SearchImage")

    image_source = "network" if bool(generateWithImages) else "preset"
    metadata = {
        "user_id": user_id,
        "search_engine": search_engine,
        "language": language,
        "generate_with_images": bool(generateWithImages),
        "image_source": image_source,
    }
    if kb_folder_ids:
        metadata["kb_folder_ids"] = kb_folder_ids
    if kb_file_ids:
        metadata["kb_file_ids"] = kb_file_ids
    logger.info(f"前端*内容**=====>metadata数据为：{metadata}")

    last_flush = asyncio.get_event_loop().time()

    try:
        async for chunk_data in content_wrapper.generate(user_question=result, metadata=metadata):
            logger.info(f"生成正文输出的chunk_data: {chunk_data}")

            # 心跳：每10秒发一次注释，避免某些代理断连接
            now = asyncio.get_event_loop().time()
            if now - last_flush > 10:
                yield b": keep-alive\n\n"
                last_flush = now

            chunk_type = chunk_data.get("type")
            if chunk_type == "text":
                # 注意：每条 SSE 事件以空行结束
                payload = chunk_data.get("text", "")
                yield _encode_sse_data(payload)
            elif chunk_type in {"error", "final"}:
                # 返回结构化事件，便于前端/日志诊断
                payload = json.dumps(chunk_data, ensure_ascii=False)
                yield _encode_sse_data(payload)
    except asyncio.CancelledError:
        logger.info("客户端已断开 PPT 内容 SSE 连接，提前结束流")
        raise
    except Exception as e:
        logger.error("内容生成流异常: %s", e, exc_info=True)
        payload = json.dumps(
            {
                "type": "error",
                "text": f"内容生成中断：{e}",
                "author": "system",
            },
            ensure_ascii=False,
        )
        yield _encode_sse_data(payload)
    finally:
        # 显式结束信号（前端可据此收尾）
        yield b"data: [DONE]\n\n"


def _inject_markdown_after_first_h1(markdown_content: str, injection_markdown: str) -> str:
    """
    将 injection_markdown 注入到 markdown_content 的首个 H1（# ）之后。
    目的：避免 stream_content_response 内部对首个 `#` 的截取逻辑把注入内容裁掉。
    """
    src = str(markdown_content or "")
    inj = (injection_markdown or "").strip()
    if not inj:
        return src
    if inj in src:
        return src

    lines = src.splitlines()
    for idx, line in enumerate(lines):
        if line.startswith("# "):
            injection_lines = ["", *inj.splitlines(), ""]
            return "\n".join(lines[: idx + 1] + injection_lines + lines[idx + 1 :])
    return inj + "\n\n" + src


@app.post("/tools/ppt")
@app.post("/tools/aippt")
async def aippt_content(request: AipptContentRequest):
    personaldb_url = _get_personaldb_url()
    personaldb_ready = bool(personaldb_url and await _is_personaldb_ready(personaldb_url))

    markdown_content = request.content
    # 兼容旧字段名：如果 user_id 为空就用 sessionId
    user_id = str(getattr(request, "user_id", None) or getattr(request, "sessionId", None) or "").strip() or "default_user"

    resolved_kb_file_ids = _normalize_kb_file_ids(request.kb_file_ids)
    full_ids, rag_ids = _split_kb_file_ids(resolved_kb_file_ids)

    # gen: 产物全文注入（不依赖 generateFromUploadedFile 开关）
    if full_ids and personaldb_ready and COURSE_OUTPUTS_START_MARKER not in (markdown_content or ""):
        full_context = await _load_personaldb_full_text_context(
            personaldb_url,
            user_id=user_id,
            file_ids=full_ids,
        )
        injection = build_course_outputs_injection_markdown(full_context, language=request.language)
        if injection:
            markdown_content = _inject_markdown_after_first_h1(markdown_content, injection)

    generate_from_uploaded_file = bool(request.generateFromUploadedFile) and personaldb_ready
    if bool(request.generateFromUploadedFile) and not generate_from_uploaded_file:
        logger.info("personaldb 不可用或未配置，强制禁用 generateFromUploadedFile: %s", personaldb_url)

    async def event_generator():
        async for chunk in stream_content_response(
            markdown_content,
            language=request.language,
            generateFromUploadedFile=generate_from_uploaded_file,
            generateFromWebSearch=request.generateFromWebSearch,
            generateWithImages=request.generateWithImages,
            user_id=user_id,
            kb_folder_ids=request.kb_folder_ids if generate_from_uploaded_file else None,
            kb_file_ids=rag_ids if generate_from_uploaded_file else None,
        ):
            yield chunk

    # 关键：SSE 推荐这些头
    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.post("/tools/lesson_plan")
async def lesson_plan(request: LessonPlanRequest):
    """
    教案生成（SSE）：
    - 每个 section 输出一条 JSON 事件
    - 结束以 data: [DONE] 收尾
    """

    async def event_generator():
        async for chunk in stream_lesson_plan_sse(request):
            yield chunk

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


# 教案导出（docx）模板列表
LESSON_DOCX_TEMPLATES: list[dict[str, str]] = [
    {"id": "lesson_simple", "name": "简洁版", "description": "分节标题 + 列表"},
    {"id": "lesson_table", "name": "表格版", "description": "流程表格布局"},
    {"id": "lesson_jnu_form", "name": "教案表单（字段）", "description": "授课题目/授课类型/教学内容/手段与方法/作业/参考资料"},
]


@app.get("/lesson/templates")
async def get_lesson_templates():
    """
    返回教案 Word（docx）导出可选模板列表（供前端选择）。
    """
    return {"data": LESSON_DOCX_TEMPLATES}


def _lesson_safe_export_filename(title: str) -> str:
    """
    构造可用于 Content-Disposition 的 docx 文件名。
    """
    base = _kb_safe_filename(title) or "lesson_plan"
    if base.lower().endswith(".docx"):
        return base
    return f"{base}.docx"


def _build_lesson_docx_bytes(*, plan: LessonPlan, style: LessonStyle, language: str, template_id: str | None = None) -> bytes:
    """
    生成 docx bytes（python-docx）。
    - 支持按 template_id 选择不同版式（类似 PPT 模板选择）
    """
    try:
        from docx import Document
        from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, Cm
    except Exception as exc:  # pragma: no cover - 依赖缺失时给出明确错误
        raise RuntimeError(f"缺少 python-docx 依赖：{exc}") from exc

    lang = (language or "zh").strip().lower()
    want_english = lang in {"en", "english"}

    tpl = _normalize_lesson_template_id(template_id)

    def _safe_float(value: Any, fallback: float) -> float:
        try:
            return float(value)
        except Exception:
            logger.debug("将 %r 转为 float 失败，使用 fallback %s", value, fallback, exc_info=True)
            return fallback

    def _safe_text(value: Any, fallback: str) -> str:
        text = str(value or "").strip()
        return text if text else fallback

    def _set_style_font(doc: Any, style_name: str, *, size_pt: int):
        try:
            st = doc.styles[style_name]
        except Exception:
            logger.debug("获取样式 %s 失败", style_name, exc_info=True)
            return
        st.font.name = style.fontZh
        st.font.size = Pt(int(size_pt))
        try:
            st._element.rPr.rFonts.set(qn("w:eastAsia"), style.fontZh)
        except Exception:
            logger.debug("设置样式 %s 中文字体失败", style_name, exc_info=True)
            pass

    def _set_run_font(run: Any, *, size_pt: int, bold: bool = False):
        run.font.name = style.fontZh
        run.font.size = Pt(int(size_pt))
        run.bold = bool(bold)
        try:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), style.fontZh)
        except Exception:
            logger.debug("设置 run 中文字体失败", exc_info=True)
            pass

    def _add_paragraph(doc: Any, text: str, *, size_pt: int, bold: bool = False, style_name: str | None = None, alignment: Any | None = None):
        p = doc.add_paragraph()
        if style_name:
            try:
                p.style = style_name
            except Exception:
                logger.debug("设置段落样式 %s 失败", style_name, exc_info=True)
                pass
        if alignment is not None:
            try:
                p.alignment = alignment
            except Exception:
                logger.debug("设置段落对齐失败", exc_info=True)
                pass
        run = p.add_run(_safe_text(text, "N/A" if want_english else "—"))
        _set_run_font(run, size_pt=size_pt, bold=bold)
        try:
            p.paragraph_format.line_spacing = float(style.lineSpacing)
        except Exception:
            logger.debug("设置段落行距失败", exc_info=True)
            pass
        return p

    def _set_cell_text(cell: Any, text: str, *, size_pt: int, bold: bool = False, alignment: Any | None = None):
        # cell.text 会重置段落样式，这里用 paragraph/run 精细控制字体
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        p.text = ""
        if alignment is not None:
            try:
                p.alignment = alignment
            except Exception:
                logger.debug("设置单元格对齐失败", exc_info=True)
                pass
        run = p.add_run(_safe_text(text, "N/A" if want_english else "—"))
        _set_run_font(run, size_pt=size_pt, bold=bold)
        try:
            p.paragraph_format.line_spacing = float(style.lineSpacing)
        except Exception:
            logger.debug("设置单元格行距失败", exc_info=True)
            pass
        return p

    def _set_cell_paragraphs(
        cell: Any,
        paragraphs: list[tuple[str, bool]],
        *,
        size_pt: int,
        alignment: Any | None = None,
    ):
        """
        以"多段落"方式写入单元格，便于实现"标签（加粗）+ 内容（多行）"的表单样式。
        - paragraphs: [(text, bold), ...]
        """
        # cell.text 会重置段落，这里借助它清空后再逐段写入
        cell.text = ""
        for idx, (text, bold) in enumerate(paragraphs):
            p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
            p.text = ""
            if alignment is not None:
                try:
                    p.alignment = alignment
                except Exception:
                    logger.debug("设置单元格段落对齐失败", exc_info=True)
                    pass
            raw = str(text or "")
            if raw:
                run = p.add_run(raw)
                _set_run_font(run, size_pt=size_pt, bold=bool(bold))
            try:
                p.paragraph_format.line_spacing = float(style.lineSpacing)
            except Exception:
                logger.debug("设置单元格段落行距失败", exc_info=True)
                pass

    def _set_run_font(run: Any, *, size_pt: int, bold: bool = False):
        run.font.name = style.fontZh
        run.font.size = Pt(int(size_pt))
        run.bold = bool(bold)
        try:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), style.fontZh)
        except Exception:
            logger.debug("设置 run 中文字体失败", exc_info=True)
            pass

    def _add_paragraph(doc: Any, text: str, *, size_pt: int, bold: bool = False, style_name: str | None = None, alignment: Any | None = None):
        p = doc.add_paragraph()
        if style_name:
            try:
                p.style = style_name
            except Exception:
                logger.debug("设置段落样式 %s 失败", style_name, exc_info=True)
                pass
        if alignment is not None:
            try:
                p.alignment = alignment
            except Exception:
                logger.debug("设置段落对齐失败", exc_info=True)
                pass
        run = p.add_run(_safe_text(text, "N/A" if want_english else "—"))
        _set_run_font(run, size_pt=size_pt, bold=bold)
        try:
            p.paragraph_format.line_spacing = float(style.lineSpacing)
        except Exception:
            logger.debug("设置段落行距失败", exc_info=True)
            pass
        return p

    def _set_cell_text(cell: Any, text: str, *, size_pt: int, bold: bool = False, alignment: Any | None = None):
        # cell.text 会重置段落样式，这里用 paragraph/run 精细控制字体
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        p.text = ""
        if alignment is not None:
            try:
                p.alignment = alignment
            except Exception:
                logger.debug("设置单元格对齐失败", exc_info=True)
                pass
        run = p.add_run(_safe_text(text, "N/A" if want_english else "—"))
        _set_run_font(run, size_pt=size_pt, bold=bold)
        try:
            p.paragraph_format.line_spacing = float(style.lineSpacing)
        except Exception:
            logger.debug("设置单元格行距失败", exc_info=True)
            pass
        return p

    def _set_cell_paragraphs(
        cell: Any,
        paragraphs: list[tuple[str, bool]],
        *,
        size_pt: int,
        alignment: Any | None = None,
    ):
        """
        以"多段落"方式写入单元格，便于实现"标签（加粗）+ 内容（多行）"的表单样式。
        - paragraphs: [(text, bold), ...]
        """
        # cell.text 会重置段落，这里借助它清空后再逐段写入
        cell.text = ""
        for idx, (text, bold) in enumerate(paragraphs):
            p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
            p.text = ""
            if alignment is not None:
                try:
                    p.alignment = alignment
                except Exception:
                    logger.debug("设置单元格段落对齐失败", exc_info=True)
                    pass
            raw = str(text or "")
            if raw:
                run = p.add_run(raw)
                _set_run_font(run, size_pt=size_pt, bold=bool(bold))
            try:
                p.paragraph_format.line_spacing = float(style.lineSpacing)
            except Exception:
                logger.debug("设置单元格段落行距失败", exc_info=True)
                pass

    def _add_paragraph(doc: Any, text: str, *, size_pt: int, bold: bool = False, style_name: str | None = None, alignment: Any | None = None):
        p = doc.add_paragraph()
        if style_name:
            try:
                p.style = style_name
            except Exception:
                logger.debug("设置段落样式 %s 失败", style_name, exc_info=True)
                pass
        if alignment is not None:
            try:
                p.alignment = alignment
            except Exception:
                logger.debug("设置段落对齐失败", exc_info=True)
                pass
        run = p.add_run(_safe_text(text, "N/A" if want_english else "—"))
        _set_run_font(run, size_pt=size_pt, bold=bold)
        try:
            p.paragraph_format.line_spacing = float(style.lineSpacing)
        except Exception:
            logger.debug("设置段落行距失败", exc_info=True)
            pass
        return p

    def _set_cell_text(cell: Any, text: str, *, size_pt: int, bold: bool = False, alignment: Any | None = None):
        # cell.text 会重置段落样式，这里用 paragraph/run 精细控制字体
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        p.text = ""
        if alignment is not None:
            try:
                p.alignment = alignment
            except Exception:
                logger.debug("设置单元格对齐失败", exc_info=True)
                pass
        run = p.add_run(_safe_text(text, "N/A" if want_english else "—"))
        _set_run_font(run, size_pt=size_pt, bold=bold)
        try:
            p.paragraph_format.line_spacing = float(style.lineSpacing)
        except Exception:
            logger.debug("设置单元格行距失败", exc_info=True)
            pass
        return p

    def _set_cell_paragraphs(
        cell: Any,
        paragraphs: list[tuple[str, bool]],
        *,
        size_pt: int,
        alignment: Any | None = None,
    ):
        """
        以"多段落"方式写入单元格，便于实现"标签（加粗）+ 内容（多行）"的表单样式。
        - paragraphs: [(text, bold), ...]
        """
        # cell.text 会重置段落，这里借助它清空后再逐段写入
        cell.text = ""
        for idx, (text, bold) in enumerate(paragraphs):
            p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
            p.text = ""
            if alignment is not None:
                try:
                    p.alignment = alignment
                except Exception:
                    logger.debug("设置单元格段落对齐失败", exc_info=True)
                    pass
            raw = str(text or "")
            if raw:
                run = p.add_run(raw)
                _set_run_font(run, size_pt=size_pt, bold=bool(bold))
            try:
                p.paragraph_format.line_spacing = float(style.lineSpacing)
            except Exception:
                logger.debug("设置单元格段落行距失败", exc_info=True)
                pass
        if alignment is not None:
            try:
                p.alignment = alignment
            except Exception:
                logger.debug("设置段落对齐失败", exc_info=True)
                pass
        run = p.add_run(_safe_text(text, "N/A" if want_english else "—"))
        _set_run_font(run, size_pt=size_pt, bold=bold)
        try:
            p.paragraph_format.line_spacing = float(style.lineSpacing)
        except Exception:
            logger.debug("设置段落行距失败", exc_info=True)
            pass
        return p

    def _set_cell_text(cell: Any, text: str, *, size_pt: int, bold: bool = False, alignment: Any | None = None):
        # cell.text 会重置段落样式，这里用 paragraph/run 精细控制字体
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        p.text = ""
        if alignment is not None:
            try:
                p.alignment = alignment
            except Exception:
                logger.debug("设置单元格对齐失败", exc_info=True)
                pass
        run = p.add_run(_safe_text(text, "N/A" if want_english else "—"))
        _set_run_font(run, size_pt=size_pt, bold=bold)
        try:
            p.paragraph_format.line_spacing = float(style.lineSpacing)
        except Exception:
            logger.debug("设置单元格行距失败", exc_info=True)
            pass
        return p

    def _set_cell_paragraphs(
        cell: Any,
        paragraphs: list[tuple[str, bool]],
        *,
        size_pt: int,
        alignment: Any | None = None,
    ):
        """
        以"多段落"方式写入单元格，便于实现"标签（加粗）+ 内容（多行）"的表单样式。
        - paragraphs: [(text, bold), ...]
        """
        # cell.text 会重置段落，这里借助它清空后再逐段写入
        cell.text = ""
        for idx, (text, bold) in enumerate(paragraphs):
            p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
            p.text = ""
            if alignment is not None:
                try:
                    p.alignment = alignment
                except Exception:
                    logger.debug("设置单元格段落对齐失败", exc_info=True)
                    pass
            raw = str(text or "")
            if raw:
                run = p.add_run(raw)
                _set_run_font(run, size_pt=size_pt, bold=bool(bold))
            try:
                p.paragraph_format.line_spacing = float(style.lineSpacing)
            except Exception:
                logger.debug("设置单元格段落行距失败", exc_info=True)
                pass

    def _shade_cell(cell: Any, *, fill_hex: str):
        """
        给单元格加底色（用于表头/标签列）。
        - 不强依赖，失败则忽略（避免不同 docx 实现差异导致导出失败）
        """
        try:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill_hex)
            tc_pr.append(shd)
        except Exception:
            logger.debug("设置单元格底色失败", exc_info=True)
            return

    def _prepare_doc() -> Any:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(_safe_float(style.marginTopCm, 2.54))
        section.bottom_margin = Cm(_safe_float(style.marginBottomCm, 2.54))
        section.left_margin = Cm(_safe_float(style.marginLeftCm, 2.54))
        section.right_margin = Cm(_safe_float(style.marginRightCm, 2.54))

        # 常用样式字体（尽量覆盖列表/标题，避免导出后字体不一致）
        _set_style_font(doc, "Normal", size_pt=style.bodySizePt)
        _set_style_font(doc, "Title", size_pt=style.titleSizePt)
        _set_style_font(doc, "Heading 1", size_pt=style.h1SizePt)
        _set_style_font(doc, "Heading 2", size_pt=style.h2SizePt)
        _set_style_font(doc, "List Bullet", size_pt=style.bodySizePt)
        _set_style_font(doc, "List Number", size_pt=style.bodySizePt)
        _set_style_font(doc, "Table Grid", size_pt=style.bodySizePt)
        return doc

    def _build_common_sections(doc: Any):
        """
        共用分节内容（目标/材料/作业）。不同模板仅在“元信息/流程”呈现上有差异。
        """
        # Objectives
        _add_paragraph(doc, "Objectives" if want_english else "教学目标", size_pt=style.h1SizePt, bold=True, style_name="Heading 1")
        if plan.objectives:
            for item in plan.objectives:
                _add_paragraph(doc, str(item), size_pt=style.bodySizePt, style_name="List Bullet")
        else:
            _add_paragraph(
                doc,
                "No objectives provided." if want_english else "未提供教学目标。",
                size_pt=style.bodySizePt,
                style_name="List Bullet",
            )

        # Materials
        _add_paragraph(doc, "Materials" if want_english else "教学材料", size_pt=style.h1SizePt, bold=True, style_name="Heading 1")
        if plan.materials:
            for item in plan.materials:
                _add_paragraph(doc, str(item), size_pt=style.bodySizePt, style_name="List Bullet")
        else:
            _add_paragraph(
                doc,
                "No materials provided." if want_english else "未提供教学材料。",
                size_pt=style.bodySizePt,
                style_name="List Bullet",
            )

        # Homework
        _add_paragraph(doc, "Homework" if want_english else "课后作业", size_pt=style.h1SizePt, bold=True, style_name="Heading 1")
        _add_paragraph(
            doc,
            (plan.homework or "").strip() or ("No homework provided." if want_english else "未提供课后作业。"),
            size_pt=style.bodySizePt,
        )

    def _build_simple(doc: Any):
        # Title
        _add_paragraph(
            doc,
            (plan.title or "").strip() or ("Lesson Plan" if want_english else "教案"),
            size_pt=style.titleSizePt,
            bold=True,
            style_name="Title",
        )

        # Meta（单行）
        aud_label = "Audience" if want_english else "受众"
        dur_label = "Duration" if want_english else "时长"
        _add_paragraph(
            doc,
            f"{aud_label}：{(plan.targetAudience or '').strip() or ('Students' if want_english else '中学学生')} | "
            f"{dur_label}：{(plan.duration or '').strip() or ('45 min' if want_english else '45分钟')}",
            size_pt=style.bodySizePt,
        )

        # Procedure（编号列表）
        _add_paragraph(doc, "Procedure" if want_english else "教学流程", size_pt=style.h1SizePt, bold=True, style_name="Heading 1")
        if plan.procedure:
            for item in plan.procedure:
                step = (item.step or "").strip() or ("Untitled Step" if want_english else "未命名步骤")
                duration = (item.duration or "").strip() or ("N/A" if want_english else "未填写时长")
                activity = (item.activity or "").strip() or ("No activity details." if want_english else "未填写活动说明。")
                if want_english:
                    _add_paragraph(doc, f"{step} ({duration})", size_pt=style.bodySizePt, bold=True, style_name="List Number")
                else:
                    _add_paragraph(doc, f"{step}（{duration}）", size_pt=style.bodySizePt, bold=True, style_name="List Number")
                _add_paragraph(doc, activity, size_pt=style.bodySizePt)
        else:
            _add_paragraph(
                doc,
                "No procedure provided." if want_english else "未提供教学流程。",
                size_pt=style.bodySizePt,
                style_name="List Number",
            )

        _build_common_sections(doc)

    def _build_table(doc: Any):
        # Title（居中）
        _add_paragraph(
            doc,
            (plan.title or "").strip() or ("Lesson Plan" if want_english else "教案"),
            size_pt=style.titleSizePt,
            bold=True,
            style_name="Title",
            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )

        # Meta（表格：受众/时长）
        aud_label = "Audience" if want_english else "受众"
        dur_label = "Duration" if want_english else "时长"
        meta = doc.add_table(rows=1, cols=4)
        try:
            meta.style = "Table Grid"
        except Exception:
            logger.debug("设置 meta 表格样式失败", exc_info=True)
            pass
        try:
            meta.alignment = WD_TABLE_ALIGNMENT.CENTER
        except Exception:
            logger.debug("设置 meta 表格对齐失败", exc_info=True)
            pass
        for cell in meta.rows[0].cells:
            try:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            except Exception:
                logger.debug("设置 meta 单元格垂直对齐失败", exc_info=True)
                pass

        meta_labels = [aud_label, (plan.targetAudience or "").strip() or ("Students" if want_english else "中学学生"), dur_label, (plan.duration or "").strip() or ("45 min" if want_english else "45分钟")]
        for idx, cell in enumerate(meta.rows[0].cells):
            is_label = idx % 2 == 0
            if is_label:
                _shade_cell(cell, fill_hex="F3F4F6")  # 灰底
            _set_cell_text(
                cell,
                meta_labels[idx],
                size_pt=style.bodySizePt,
                bold=is_label,
                alignment=WD_PARAGRAPH_ALIGNMENT.CENTER if is_label else WD_PARAGRAPH_ALIGNMENT.LEFT,
            )

        # Procedure（表格：环节/时长/活动）
        _add_paragraph(doc, "Procedure" if want_english else "教学流程", size_pt=style.h1SizePt, bold=True, style_name="Heading 1")
        items = plan.procedure or []
        rows = (len(items) if items else 1) + 1  # + 表头
        table = doc.add_table(rows=rows, cols=3)
        try:
            table.style = "Table Grid"
        except Exception:
            logger.debug("设置 procedure 表格样式失败", exc_info=True)
            pass
        try:
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
        except Exception:
            logger.debug("设置 procedure 表格对齐失败", exc_info=True)
            pass
        try:
            table.autofit = False
        except Exception:
            logger.debug("设置 procedure 表格 autofit 失败", exc_info=True)
            pass

        # 尝试设置列宽（不保证所有 Word 版本完全一致，但可提升可读性）
        try:
            table.columns[0].width = Cm(4)
            table.columns[1].width = Cm(3)
            table.columns[2].width = Cm(10)
        except Exception:
            logger.debug("设置 procedure 表格列宽失败", exc_info=True)
            pass

        headers = ["Step" if want_english else "环节", "Duration" if want_english else "时长", "Activity" if want_english else "活动"]
        header_row = table.rows[0]
        for i, cell in enumerate(header_row.cells):
            _shade_cell(cell, fill_hex="E5E7EB")  # 深一点的灰
            try:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            except Exception:
                logger.debug("设置 procedure 表头垂直对齐失败", exc_info=True)
                pass
            _set_cell_text(cell, headers[i], size_pt=style.bodySizePt, bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)

        if items:
            for idx, item in enumerate(items, start=1):
                step = (item.step or "").strip() or ("Untitled Step" if want_english else "未命名步骤")
                duration = (item.duration or "").strip() or ("N/A" if want_english else "未填写时长")
                activity = (item.activity or "").strip() or ("No activity details." if want_english else "未填写活动说明。")
                row = table.rows[idx]
                for cell in row.cells:
                    try:
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                    except Exception:
                        logger.debug("设置 procedure 单元格垂直对齐失败", exc_info=True)
                        pass
                _set_cell_text(row.cells[0], f"{idx}. {step}", size_pt=style.bodySizePt, bold=True)
                _set_cell_text(row.cells[1], duration, size_pt=style.bodySizePt, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
                _set_cell_text(row.cells[2], activity, size_pt=style.bodySizePt)
        else:
            row = table.rows[1]
            _set_cell_text(row.cells[0], "—", size_pt=style.bodySizePt)
            _set_cell_text(row.cells[1], "—", size_pt=style.bodySizePt)
            _set_cell_text(row.cells[2], "No procedure provided." if want_english else "未提供教学流程。", size_pt=style.bodySizePt)

        _build_common_sections(doc)

    def _build_jnu_form(doc: Any):
        """
        “字段表单”模板（参考用户提供的教案表格）：
        - 授课题目、授课类型、授课时间（日期/周次节次）
        - 教学内容（基本内容/重点/难点）
        - 教学手段与方法、作业、参考资料
        """

        # 表头/字段标签（根据语言做最小翻译；该模板主要面向中文）
        if want_english:
            topic_label = "Teaching Topic (chapter/theme):"
            lesson_type_label = "Lesson Type"
            lesson_time_label = "Lesson Time"
            date_placeholder = "YYYY  MM  DD"
            period_placeholder = "Week __  Day __  Period __"
            content_label = "Teaching Content (basics / key / difficult points):"
            methods_label = "Teaching Tools and Methods:"
            homework_label = "Questions / Discussion / Homework:"
            refs_label = "References (books, papers, etc.):"
            basics_label = "Basics:"
            key_label = "Key Points:"
            diff_label = "Difficult Points:"
            tools_label = "Tools:"
            method_label = "Methods:"
        else:
            topic_label = "授课题目（教学章节或主题）："
            lesson_type_label = "授课类型"
            lesson_time_label = "授课时间"
            date_placeholder = "年   月   日"
            period_placeholder = "第   周星期     第   节"
            content_label = "教学内容（包括基本内容、重点、难点三部分）："
            methods_label = "教学手段与方法："
            homework_label = "思考题、讨论题或作业："
            refs_label = "参考资料（包括辅助教材、参考书、文献等）："
            basics_label = "基本内容："
            key_label = "重点："
            diff_label = "难点："
            tools_label = "教学手段："
            method_label = "教学方法："

        def _infer_lesson_type() -> str:
            """
            尝试根据流程/活动粗略推断授课类型。
            - 仅用于填表默认值，避免空白
            """
            haystack = " ".join(
                [
                    f"{(p.step or '').strip()} {(p.activity or '').strip()}".strip()
                    for p in (plan.procedure or [])
                ]
            )
            if want_english:
                if "实验" in haystack or "experiment" in haystack.lower():
                    return "Lab"
                if "讨论" in haystack or "discussion" in haystack.lower():
                    return "Discussion"
                if "实践" in haystack or "实习" in haystack or "见习" in haystack:
                    return "Practice"
                return "Lecture"
            else:
                if "实验" in haystack:
                    return "实验课"
                if "讨论" in haystack:
                    return "讨论课"
                if "实践" in haystack or "实习" in haystack or "见习" in haystack:
                    return "实践课"
                return "理论课"

        def _build_basic_content_lines() -> list[str]:
            if not plan.procedure:
                return ["—" if not want_english else "N/A"]
            lines: list[str] = []
            for idx, item in enumerate(plan.procedure, start=1):
                step = (item.step or "").strip() or ("Untitled Step" if want_english else "未命名步骤")
                duration = (item.duration or "").strip()
                activity = (item.activity or "").strip()
                if duration and activity:
                    lines.append(f"{idx}. {step} ({duration}) {activity}" if want_english else f"{idx}. {step}（{duration}）{activity}")
                elif duration:
                    lines.append(f"{idx}. {step} ({duration})" if want_english else f"{idx}. {step}（{duration}）")
                elif activity:
                    lines.append(f"{idx}. {step}: {activity}" if want_english else f"{idx}. {step}：{activity}")
                else:
                    lines.append(f"{idx}. {step}")
            return lines

        def _build_methods() -> tuple[str, str]:
            # 手段：优先用 materials 列表
            tools = "、".join([str(x).strip() for x in (plan.materials or []) if str(x).strip()]) if not want_english else ", ".join([str(x).strip() for x in (plan.materials or []) if str(x).strip()])
            tools = tools or ("—" if not want_english else "N/A")

            # 方法：用简单规则从流程/活动中抽取关键词
            haystack = " ".join(
                [
                    f"{(p.step or '').strip()} {(p.activity or '').strip()}".strip()
                    for p in (plan.procedure or [])
                ]
            )
            method_candidates: list[str] = []
            if want_english:
                method_candidates.append("Lecture")
                if "讨论" in haystack or "discussion" in haystack.lower():
                    method_candidates.append("Discussion")
                if "示范" in haystack or "demo" in haystack.lower() or "demonstr" in haystack.lower():
                    method_candidates.append("Demonstration")
                if "练习" in haystack or "exercise" in haystack.lower() or "practice" in haystack.lower():
                    method_candidates.append("Practice")
                if "小组" in haystack or "合作" in haystack or "group" in haystack.lower():
                    method_candidates.append("Group work")
                if "提问" in haystack or "问答" in haystack or "q&a" in haystack.lower() or "question" in haystack.lower():
                    method_candidates.append("Q&A")
                # 去重保持顺序
                uniq: list[str] = []
                for m in method_candidates:
                    if m not in uniq:
                        uniq.append(m)
                methods = ", ".join(uniq) if uniq else "N/A"
            else:
                method_candidates.append("讲授")
                if "讨论" in haystack:
                    method_candidates.append("讨论")
                if "示范" in haystack:
                    method_candidates.append("示范")
                if "练习" in haystack or "训练" in haystack:
                    method_candidates.append("练习")
                if "小组" in haystack or "合作" in haystack:
                    method_candidates.append("小组合作")
                if "提问" in haystack or "问答" in haystack:
                    method_candidates.append("提问")
                uniq = []
                for m in method_candidates:
                    if m not in uniq:
                        uniq.append(m)
                methods = "、".join(uniq) if uniq else "—"

            return tools, methods

        def _build_refs_lines() -> list[str]:
            # 参考资料：优先从 materials 中挑选“看起来像参考书/教材/文献”的项；否则留空
            keywords = ("教材", "课本", "参考", "文献", "论文", "书", "ISBN")
            refs = [str(x).strip() for x in (plan.materials or []) if str(x).strip() and any(k in str(x) for k in keywords)]
            if not refs:
                return ["—" if not want_english else "N/A"]
            return [f"- {r}" for r in refs] if not want_english else [f"- {r}" for r in refs]

        lesson_type_value = _infer_lesson_type()

        # 7x3 表格，与用户示例一致：左侧（题目）跨 3 行，授课时间标签跨 2 行，后续 4 行各自跨 3 列
        table = doc.add_table(rows=7, cols=3)
        try:
            table.style = "Table Grid"
        except Exception:
            logger.debug("设置 JNU 表格样式失败", exc_info=True)
            pass
        try:
            table.autofit = False
        except Exception:
            logger.debug("设置 JNU 表格 autofit 失败", exc_info=True)
            pass

        # 尝试设置列宽：A4 扣边距后约 15.9cm，可按需微调
        try:
            table.columns[0].width = Cm(9.0)
            table.columns[1].width = Cm(2.6)
            table.columns[2].width = Cm(4.3)
        except Exception:
            logger.debug("设置 JNU 表格列宽失败", exc_info=True)
            pass

        # 合并单元格（python-docx 用矩形区域合并）
        topic_cell = table.cell(0, 0).merge(table.cell(2, 0))
        time_label_cell = table.cell(1, 1).merge(table.cell(2, 1))
        for r in range(3, 7):
            table.cell(r, 0).merge(table.cell(r, 2))

        # 基础对齐
        for r in range(7):
            for c in range(3):
                cell = table.cell(r, c)
                try:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                except Exception:
                    logger.debug("设置 JNU 单元格垂直对齐失败", exc_info=True)
                    pass

        # Row 0-2: 授课题目 / 类型 / 时间
        _set_cell_paragraphs(
            topic_cell,
            [
                (topic_label, True),
                (_safe_text(plan.title, "Lesson Plan" if want_english else "教案"), False),
            ],
            size_pt=style.bodySizePt,
        )

        _shade_cell(table.cell(0, 1), fill_hex="F3F4F6")
        _set_cell_text(table.cell(0, 1), lesson_type_label, size_pt=style.bodySizePt, bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
        _set_cell_text(table.cell(0, 2), lesson_type_value, size_pt=style.bodySizePt)

        _shade_cell(time_label_cell, fill_hex="F3F4F6")
        _set_cell_text(time_label_cell, lesson_time_label, size_pt=style.bodySizePt, bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
        _set_cell_text(table.cell(1, 2), date_placeholder, size_pt=style.bodySizePt)
        _set_cell_text(table.cell(2, 2), period_placeholder, size_pt=style.bodySizePt)

        # Row 3: 教学内容（基本/重点/难点）
        basic_lines = _build_basic_content_lines()

        obj_list = [str(x).strip() for x in (plan.objectives or []) if str(x).strip()]
        key_objs: list[str] = []
        diff_text = "N/A" if want_english else "—"
        if obj_list:
            last = obj_list[-1]
            is_diff = bool(re.match(r"^difficulty\s*[:：]", last, re.IGNORECASE)) if want_english else bool(re.match(r"^难点\s*[:：]", last))
            if is_diff:
                key_objs = obj_list[:-1]
                diff_text = re.sub(r"^difficulty\s*[:：]\s*", "", last, flags=re.IGNORECASE) if want_english else re.sub(r"^难点\s*[:：]\s*", "", last)
            elif len(obj_list) >= 2:
                key_objs = obj_list[:-1]
                diff_text = last
            else:
                key_objs = obj_list

        key_lines = [f"- {x}" for x in key_objs if str(x).strip()] or ["—" if not want_english else "N/A"]
        diff_lines = [(diff_text or ("N/A" if want_english else "—"), False)]
        content_cell = table.cell(3, 0)
        paragraphs: list[tuple[str, bool]] = [(content_label, True), (basics_label, True)]
        paragraphs.extend([(line, False) for line in basic_lines])
        paragraphs.append(("", False))
        paragraphs.append((key_label, True))
        paragraphs.extend([(line, False) for line in key_lines])
        paragraphs.append(("", False))
        paragraphs.append((diff_label, True))
        paragraphs.extend(diff_lines)
        _set_cell_paragraphs(content_cell, paragraphs, size_pt=style.bodySizePt)

        # Row 4: 教学手段与方法
        tools_value, methods_value = _build_methods()
        methods_cell = table.cell(4, 0)
        _set_cell_paragraphs(
            methods_cell,
            [
                (methods_label, True),
                (f"{tools_label} {tools_value}", False),
                (f"{method_label} {methods_value}", False),
            ],
            size_pt=style.bodySizePt,
        )

        # Row 5: 作业/讨论题
        _set_cell_paragraphs(
            table.cell(5, 0),
            [
                (homework_label, True),
                (_safe_text(plan.homework, "—" if not want_english else "N/A"), False),
            ],
            size_pt=style.bodySizePt,
        )

        # Row 6: 参考资料
        ref_lines = _build_refs_lines()
        _set_cell_paragraphs(
            table.cell(6, 0),
            [(refs_label, True)] + [(line, False) for line in ref_lines],
            size_pt=style.bodySizePt,
        )

    doc = _prepare_doc()
    if tpl == "lesson_simple":
        _build_simple(doc)
    elif tpl == "lesson_table":
        _build_table(doc)
    elif tpl == "lesson_jnu_form":
        _build_jnu_form(doc)
    else:
        raise ValueError(f"未知教案模板：{tpl}")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


@app.post("/lesson/export/docx")
async def lesson_export_docx(request: LessonExportDocxRequest):
    """
    导出教案为标准 .docx（附件下载）。
    """
    plan = request.lessonPlan
    style = request.style or LessonStyle()
    language = request.language or "zh"
    template_id = (request.templateId or "").strip() or None

    try:
        content = _build_lesson_docx_bytes(plan=plan, style=style, language=language, template_id=template_id)
    except ValueError as exc:
        # 模板不合法属于 4xx（调用方参数错误），不要吞成 500
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        logger.error("lesson_export_docx 失败: %s", exc, exc_info=True)
        raise HTTPException(status_code=500, detail=str(exc)) from exc

    filename = _lesson_safe_export_filename(plan.title)
    encoded = quote(filename)

    artifact_id: str | None = None
    if bool(request.persist):
        persist_user_id = str(request.userId or "").strip()
        persist_material_id = str(request.materialId or "").strip()
        if persist_user_id and persist_material_id:
            try:
                meta = _save_artifact_bytes(
                    user_id=persist_user_id,
                    material_id=persist_material_id,
                    kind="docx",
                    file_bytes=content,
                    file_name=filename,
                )
                artifact_id = str(meta.get("artifact_id") or "").strip() or None
            except Exception as exc:
                logger.error("lesson_export_docx 持久化 artifacts 失败: %s", exc, exc_info=True)
        else:
            logger.info("lesson_export_docx persist=true 但缺少 userId/materialId，跳过持久化")

    headers = {
        "Content-Disposition": f"attachment; filename*=UTF-8''{encoded}",
        "Cache-Control": "no-store",
    }
    if artifact_id:
        headers["X-TeachDo-Artifact-Id"] = artifact_id

    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@app.post("/tools/assistant_chat")
async def assistant_chat(request: AssistantChatRequest):
    """
    助教对话（SSE）：
    - 历史消息由前端维护，每次请求透传 messages
    - 可选透传 kb_file_ids，用 personaldb 检索片段增强回答（RAG）
    - 不做会话持久化，提供 “清除上下文” 由前端实现（清空 messages）
    """
    if not request.messages:
        raise HTTPException(status_code=400, detail="messages 不能为空")

    last_user_message = _pick_last_user_message(request.messages)
    if not last_user_message:
        raise HTTPException(status_code=400, detail="messages 中缺少 user 消息")

    personaldb_url = _get_personaldb_url()
    resolved_kb_file_ids = _normalize_kb_file_ids(request.kb_file_ids)
    full_context = ""
    kb_context = ""
    if resolved_kb_file_ids and personaldb_url and await _is_personaldb_ready(personaldb_url):
        full_context, kb_context = await _build_personaldb_kb_contexts(
            personaldb_url,
            user_id=str(request.user_id or "default_user"),
            query=last_user_message,
            kb_file_ids=resolved_kb_file_ids,
            rag_topk=5,
        )

    system_prompt = _build_assistant_system_prompt(
        material=request.material,
        full_context=full_context,
        kb_context=kb_context,
        language=request.language,
    )

    llm_messages: list[dict[str, str]] = [{"role": "system", "content": system_prompt}]
    for msg in request.messages:
        content = (msg.content or "").strip()
        if not content:
            continue
        llm_messages.append({"role": msg.role, "content": content})

    try:
        llm_settings = _get_assistant_llm_settings()
    except Exception as exc:
        logger.error("助教模型配置错误: %s", exc, exc_info=True)
        raise HTTPException(status_code=500, detail=str(exc)) from exc

    async def event_generator():
        async for chunk in stream_assistant_sse(
            llm_settings=llm_settings,
            messages=llm_messages,
        ):
            yield chunk

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.post("/kb/upload")
async def kb_upload(
    user_id: str = Form(...),
    folder_id: int = Form(0),
    file_id: str | None = Form(None),
    file_type: str | None = Form(None),
    file: UploadFile = File(...),
):
    """
    KB BFF：上传素材并向量化（转发到 personaldb /upload/）。
    - 前端统一访问 /api/kb/upload（Vite proxy 去掉 /api）
    """
    personaldb_url = _get_personaldb_url()
    if not personaldb_url:
        return _kb_error("KB_NOT_CONFIGURED", "PERSONAL_DB 未配置", status_code=500)

    if not file:
        return _kb_error("KB_FILE_REQUIRED", "缺少文件", status_code=400)

    resolved_file_type = (file_type or "").strip() or None
    if not resolved_file_type and file.filename and "." in file.filename:
        resolved_file_type = file.filename.rsplit(".", 1)[-1]

    resolved_file_id = (file_id or "").strip() or None
    if not resolved_file_id:
        epoch_ms = int(time.time() * 1000)
        resolved_file_id = f"upload:{user_id}:{epoch_ms}:{random.randint(0, 999):03d}"

    file_bytes = await file.read()
    if not file_bytes:
        return _kb_error("KB_EMPTY_FILE", "文件内容为空", status_code=400)
    file_size = len(file_bytes)

    data = {
        "userId": str(user_id),
        "fileId": str(resolved_file_id),
        "folderId": str(folder_id),
    }
    if resolved_file_type:
        data["fileType"] = str(resolved_file_type)

    files_payload = {
        "file": (
            file.filename or "uploaded_file",
            file_bytes,
            file.content_type or "application/octet-stream",
        )
    }

    upload_url = f"{personaldb_url}/upload/"

    async with httpx.AsyncClient(trust_env=False, timeout=httpx.Timeout(360.0)) as client:
        try:
            resp = await client.post(upload_url, data=data, files=files_payload)
            if resp.status_code >= 400:
                logger.info("[personaldb %s] %s", resp.status_code, resp.text)
                return _kb_error("KB_UPLOAD_FAILED", resp.text, status_code=resp.status_code)
            result = resp.json()
        except Exception as exc:
            logger.error("kb_upload 调用 personaldb 失败: %s", exc, exc_info=True)
            return _kb_error("KB_UPLOAD_FAILED", f"personaldb 调用失败: {exc}", status_code=502)

    # 不向前端返回 markdown_content（可能很大）
    return _kb_ok(
        {
            "user_id": str(user_id),
            "file_id": str(resolved_file_id),
            "file_name": file.filename or result.get("file_name") or "uploaded_file",
            "file_type": resolved_file_type or result.get("fileType") or "unknown",
            "file_size": int(file_size),
            "folder_id": int(folder_id),
            "status": "ready",
        }
    )


@app.get("/kb/files/{user_id}")
async def kb_list_files(user_id: str, folder_id: int | None = Query(None)):
    """
    KB BFF：列出知识库文件（转发 personaldb GET /files/{user_id}）。
    """
    personaldb_url = _get_personaldb_url()
    if not personaldb_url:
        return _kb_error("KB_NOT_CONFIGURED", "PERSONAL_DB 未配置", status_code=500)

    url = f"{personaldb_url}/files/{user_id}"
    async with httpx.AsyncClient(trust_env=False, timeout=httpx.Timeout(10.0)) as client:
        try:
            resp = await client.get(url)
            if resp.status_code >= 400:
                logger.info("[personaldb %s] %s", resp.status_code, resp.text)
                return _kb_error("KB_LIST_FAILED", resp.text, status_code=resp.status_code)
            files = resp.json()
        except Exception as exc:
            logger.error("kb_list_files 调用 personaldb 失败: %s", exc, exc_info=True)
            return _kb_error("KB_LIST_FAILED", f"personaldb 调用失败: {exc}", status_code=502)

    if not isinstance(files, list):
        return _kb_error("KB_LIST_FAILED", "personaldb 返回格式非法（期望 list）", status_code=502)

    normalized = []
    for item in files:
        if not isinstance(item, dict):
            continue
        try:
            fid = str(item.get("file_id") or item.get("fileId") or "")
            if not fid:
                continue
            one_folder_id = item.get("folder_id") if item.get("folder_id") is not None else item.get("folderId")
            one_folder_id_int = int(one_folder_id) if one_folder_id is not None else 0
            one_file_size = item.get("file_size") if item.get("file_size") is not None else item.get("fileSize")
            try:
                one_file_size_int = int(one_file_size) if one_file_size is not None else 0
                if one_file_size_int < 0:
                    one_file_size_int = 0
            except Exception:
                logger.debug("忽略非法 file_size: %r", one_file_size, exc_info=True)
                one_file_size_int = 0

            raw_created_at = item.get("created_at") if item.get("created_at") is not None else item.get("createdAt")
            created_at_ms: int | None = None
            if raw_created_at is not None:
                try:
                    created_at_ms = int(raw_created_at)
                    if created_at_ms > 0 and created_at_ms < 1_000_000_000_000:
                        created_at_ms *= 1000
                except Exception:
                    logger.debug("忽略非法 created_at: %r", raw_created_at, exc_info=True)
                    created_at_ms = None

            raw_source_type = item.get("source_type") if item.get("source_type") is not None else item.get("sourceType")
            source_type = str(raw_source_type).strip().lower() if raw_source_type is not None else ""
            if source_type not in {"upload", "material"}:
                source_type = ""

            source_material_id = (
                str(
                    item.get("source_material_id")
                    if item.get("source_material_id") is not None
                    else item.get("sourceMaterialId")
                    or ""
                ).strip()
            )
            source_material_title = (
                str(
                    item.get("source_material_title")
                    if item.get("source_material_title") is not None
                    else item.get("sourceMaterialTitle")
                    or ""
                ).strip()
            )
            if folder_id is not None and int(folder_id) != one_folder_id_int:
                continue
            normalized.append(
                {
                    "user_id": str(user_id),
                    "file_id": fid,
                    "file_name": item.get("file_name") or item.get("fileName") or "",
                    "file_type": item.get("file_type") or item.get("fileType") or "",
                    "file_size": one_file_size_int,
                    "folder_id": one_folder_id_int,
                    **({"created_at": created_at_ms} if created_at_ms is not None else {}),
                    **({"source_type": source_type} if source_type else {}),
                    **({"source_material_id": source_material_id} if source_material_id else {}),
                    **({"source_material_title": source_material_title} if source_material_title else {}),
                }
            )
        except Exception:
            logger.warning("解析 KB 文件项失败，跳过该项", exc_info=True)
            continue

    return _kb_ok(normalized)


@app.get("/kb/files/{user_id}/{file_id}/export")
async def kb_export_file(user_id: str, file_id: str):
    """
    KB BFF：导出知识库文件内容（Markdown/纯文本）。

    - 转发 personaldb GET /files/{user_id}/{file_id}/content
    - 以 attachment 形式返回，便于前端下载保存
    """
    personaldb_url = _get_personaldb_url()
    if not personaldb_url:
        return _kb_error("KB_NOT_CONFIGURED", "PERSONAL_DB 未配置", status_code=500)

    url = f"{personaldb_url}/files/{user_id}/{file_id}/content"
    async with httpx.AsyncClient(trust_env=False, timeout=httpx.Timeout(20.0)) as client:
        try:
            resp = await client.get(url)
            if resp.status_code == 404:
                return _kb_error("KB_FILE_NOT_FOUND", "文件不存在", status_code=404)
            if resp.status_code >= 400:
                logger.info("[personaldb %s] %s", resp.status_code, resp.text)
                return _kb_error("KB_EXPORT_FAILED", resp.text, status_code=resp.status_code)
            payload = resp.json()
        except Exception as exc:
            logger.error("kb_export_file 调用 personaldb 失败: %s", exc, exc_info=True)
            return _kb_error("KB_EXPORT_FAILED", f"personaldb 调用失败: {exc}", status_code=502)

    content = payload.get("content") if isinstance(payload, dict) else None
    if not isinstance(content, str):
        return _kb_error("KB_EXPORT_FAILED", "personaldb 返回格式非法（缺少 content）", status_code=502)

    file_name = payload.get("file_name") if isinstance(payload, dict) else ""
    file_type = payload.get("file_type") if isinstance(payload, dict) else ""
    export_name = _kb_build_export_filename(str(file_name or ""), str(file_type or ""), file_id)
    encoded = quote(export_name)

    return Response(
        content=content.encode("utf-8"),
        media_type="text/markdown; charset=utf-8",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded}",
            "Cache-Control": "no-store",
        },
    )


@app.get("/artifacts/{user_id}/{material_id}")
async def list_artifacts(user_id: str, material_id: str):
    """
    列出指定课程的导出产物（PPTX/DOCX 等）。
    """
    material_dir = _get_artifact_material_dir(user_id=str(user_id), material_id=str(material_id))
    items = _load_artifact_index(material_dir)

    def _created_at(it: dict[str, Any]) -> int:
        raw = it.get("created_at")
        try:
            return int(raw)
        except Exception:
            logger.debug("忽略非法 artifact created_at: %r", raw, exc_info=True)
            return 0

    items_sorted = sorted(items, key=_created_at, reverse=True)
    return _kb_ok([_artifact_public_meta(it) for it in items_sorted])


@app.post("/artifacts/{user_id}/{material_id}")
async def upload_artifact(
    user_id: str,
    material_id: str,
    kind: str = Form(...),
    file: UploadFile = File(...),
):
    """
    上传一个导出产物文件（multipart/form-data）。
    fields:
    - kind: pptx | docx
    - file: 二进制文件
    """
    normalized_kind = _normalize_artifact_kind(kind)
    if not normalized_kind:
        return _kb_error("ARTIFACT_KIND_INVALID", "kind 必须是 pptx 或 docx", status_code=400)

    file_bytes = await file.read()
    if not file_bytes:
        return _kb_error("ARTIFACT_FILE_EMPTY", "文件内容为空", status_code=400)

    original_name = file.filename or f"{normalized_kind}.{normalized_kind}"
    try:
        meta = _save_artifact_bytes(
            user_id=str(user_id),
            material_id=str(material_id),
            kind=normalized_kind,
            file_bytes=file_bytes,
            file_name=original_name,
        )
    except Exception as exc:
        logger.error("upload_artifact 保存失败: %s", exc, exc_info=True)
        return _kb_error("ARTIFACT_UPLOAD_FAILED", str(exc), status_code=500)

    return _kb_ok(meta)


@app.get("/artifacts/{user_id}/{material_id}/{artifact_id}")
async def download_artifact(user_id: str, material_id: str, artifact_id: str):
    """
    下载一个导出产物文件（attachment）。
    """
    material_dir = _get_artifact_material_dir(user_id=str(user_id), material_id=str(material_id))
    items = _load_artifact_index(material_dir)

    target: dict[str, Any] | None = None
    for it in items:
        if str(it.get("artifact_id") or "") == str(artifact_id):
            target = it
            break
    if not target:
        raise HTTPException(status_code=404, detail="artifact 不存在")

    stored_name = str(target.get("stored_name") or "").strip()
    if not stored_name:
        raise HTTPException(status_code=404, detail="artifact 索引损坏（缺少 stored_name）")
    file_path = material_dir / stored_name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="artifact 文件不存在")

    file_name = str(target.get("file_name") or f"{artifact_id}").strip() or f"{artifact_id}"
    encoded = quote(_kb_safe_filename(file_name))
    media_type = _artifact_media_type(str(target.get("kind") or ""))
    return FileResponse(
        file_path,
        media_type=media_type,
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded}",
            "Cache-Control": "no-store",
        },
    )


@app.delete("/artifacts/{user_id}/{material_id}/{artifact_id}")
async def delete_artifact(user_id: str, material_id: str, artifact_id: str):
    """
    删除一个导出产物文件。
    """
    material_dir = _get_artifact_material_dir(user_id=str(user_id), material_id=str(material_id))
    items = _load_artifact_index(material_dir)

    kept: list[dict[str, Any]] = []
    target: dict[str, Any] | None = None
    for it in items:
        if str(it.get("artifact_id") or "") == str(artifact_id) and target is None:
            target = it
            continue
        kept.append(it)

    if not target:
        raise HTTPException(status_code=404, detail="artifact 不存在")

    stored_name = str(target.get("stored_name") or "").strip()
    if stored_name:
        try:
            (material_dir / stored_name).unlink()
        except FileNotFoundError:
            pass
        except Exception as exc:
            logger.info("delete_artifact 删除文件失败：%s", exc)

    _write_artifact_index(material_dir, kept)
    return _kb_ok({"artifact_id": str(artifact_id), "deleted": True})


class KbVectorizeTextRequest(BaseModel):
    user_id: str
    file_id: str
    file_name: str
    content: str
    file_type: str = "md"
    folder_id: int = 1
    # KB 元数据（可选）：用于前端展示“时间 + 来源”，不参与检索
    created_at: int | None = None
    source_type: str | None = None
    source_material_id: str | None = None
    source_material_title: str | None = None


@app.post("/kb/vectorize/text")
async def kb_vectorize_text(request: KbVectorizeTextRequest):
    """
    KB BFF：把文本写入 KB 索引（转发 personaldb POST /vectorize/text）。
    """
    personaldb_url = _get_personaldb_url()
    if not personaldb_url:
        return _kb_error("KB_NOT_CONFIGURED", "PERSONAL_DB 未配置", status_code=500)

    if not request.content.strip():
        return _kb_error("KB_CONTENT_REQUIRED", "content 不能为空", status_code=400)

    payload = {
        "userId": request.user_id,
        "fileId": request.file_id,
        "fileName": request.file_name,
        "fileType": request.file_type,
        "folderId": request.folder_id,
        "content": request.content,
        "url": "",
    }
    if request.created_at is not None:
        try:
            payload["createdAt"] = int(request.created_at)
        except Exception:
            logger.debug("忽略非法 created_at: %r", request.created_at, exc_info=True)
            pass
    if request.source_type:
        payload["sourceType"] = str(request.source_type)
    if request.source_material_id:
        payload["sourceMaterialId"] = str(request.source_material_id)
    if request.source_material_title:
        payload["sourceMaterialTitle"] = str(request.source_material_title)

    url = f"{personaldb_url}/vectorize/text"
    async with httpx.AsyncClient(trust_env=False, timeout=httpx.Timeout(60.0)) as client:
        try:
            resp = await client.post(url, json=payload)
            if resp.status_code >= 400:
                logger.info("[personaldb %s] %s", resp.status_code, resp.text)
                return _kb_error("KB_VECTORIZE_FAILED", resp.text, status_code=resp.status_code)
        except Exception as exc:
            logger.error("kb_vectorize_text 调用 personaldb 失败: %s", exc, exc_info=True)
            return _kb_error("KB_VECTORIZE_FAILED", f"personaldb 调用失败: {exc}", status_code=502)

    return _kb_ok(True)


@app.delete("/kb/files/{user_id}/{file_id}")
async def kb_delete_file(user_id: str, file_id: str):
    """
    KB BFF：删除知识库文件向量（转发 personaldb DELETE /files/{user_id}/{file_id}）。
    """
    personaldb_url = _get_personaldb_url()
    if not personaldb_url:
        return _kb_error("KB_NOT_CONFIGURED", "PERSONAL_DB 未配置", status_code=500)

    url = f"{personaldb_url}/files/{user_id}/{file_id}"
    async with httpx.AsyncClient(trust_env=False, timeout=httpx.Timeout(10.0)) as client:
        try:
            resp = await client.delete(url)
            if resp.status_code >= 400:
                logger.info("[personaldb %s] %s", resp.status_code, resp.text)
                return _kb_error("KB_DELETE_FAILED", resp.text, status_code=resp.status_code)
        except Exception as exc:
            logger.error("kb_delete_file 调用 personaldb 失败: %s", exc, exc_info=True)
            return _kb_error("KB_DELETE_FAILED", f"personaldb 调用失败: {exc}", status_code=502)

    return _kb_ok(True)

@app.get("/data/{filename}")
async def get_data(filename: str):
    resolved = resolve_safe_static_file(TEMPLATE_DIR, filename)
    if not resolved:
        raise HTTPException(status_code=404, detail="not found")
    return FileResponse(str(resolved))

@app.get("/templates")
async def get_templates():
    templates = [
        { "name": "红色通用", "id": "template_1", "cover": "/api/data/template_1.jpg" },
        { "name": "蓝色通用", "id": "template_2", "cover": "/api/data/template_2.jpg" },
        { "name": "紫色通用", "id": "template_3", "cover": "/api/data/template_3.jpg" },
        { "name": "莫兰迪配色", "id": "template_4", "cover": "/api/data/template_4.jpg" },
        # { "name": "图表", "id": "template_6", "cover": "/api/data/template_6.jpg" },
    ]

    return {"data": templates}


@app.get("/files/{user_id}")
async def list_user_files(user_id: int):
    """
    列出指定用户的所有文件信息
    """
    personaldb_api_url = os.environ["PERSONAL_DB"]
    url = f"{personaldb_api_url}/files/{user_id}"

    # 内部服务调用（personaldb）不应受系统代理环境变量影响
    async with httpx.AsyncClient(trust_env=False) as client:
        try:
            response = await client.get(url)
            response.raise_for_status()
            return response.json()
        except httpx.RequestError as exc:
            raise HTTPException(status_code=500, detail=f"Error connecting to personaldb: {exc}")
        except httpx.HTTPStatusError as exc:
            # 转发下游服务的错误
            raise HTTPException(status_code=exc.response.status_code, detail=exc.response.text)


@app.get("/proxy")
async def proxy(request: Request, url: str = Query(..., description="Target absolute URL")):
    """
    透明代理上游资源，转发部分请求头，透传关键响应头，并允许前端同源访问。
    适合图片/音视频等二进制内容。
    """
    HEADERS_TO_FORWARD = {"Range", "User-Agent"}  # 需要时可扩展
    HEADERS_TO_COPY = {
        "Content-Type",
        "Content-Length",
        "Content-Disposition",
        "Accept-Ranges",
        "ETag",
        "Last-Modified",
        "Cache-Control",
        "Expires",
    }
    forward_headers = {}
    for h in HEADERS_TO_FORWARD:
        v = request.headers.get(h)
        if v:
            forward_headers[h] = v

    try:
        client, upstream = await _open_validated_proxy_stream(url, headers=forward_headers)
    except UrlAccessError as exc:
        raise HTTPException(status_code=exc.status_code, detail=str(exc)) from exc
    except httpx.RequestError as exc:
        raise HTTPException(status_code=502, detail=f"Upstream fetch error: {exc!s}") from exc

    if upstream.status_code >= 400:
        await upstream.aclose()
        await client.aclose()
        raise HTTPException(status_code=upstream.status_code, detail="Upstream error")

    max_bytes = get_proxy_max_bytes()
    upstream_content_length = upstream.headers.get("Content-Length")
    if upstream_content_length:
        try:
            if int(upstream_content_length) > max_bytes:
                await upstream.aclose()
                await client.aclose()
                raise HTTPException(status_code=413, detail="上游资源过大")
        except ValueError:
            # 非法 Content-Length 不阻断，交给流式计数兜底
            pass

    headers = {}
    for h in HEADERS_TO_COPY:
        if h in upstream.headers:
            headers[h] = upstream.headers[h]

    # 给静态资源加简单缓存（按需调整）
    headers.setdefault("Cache-Control", "public, max-age=86400")

    async def _limited_iter_bytes():
        total = 0
        async for chunk in upstream.aiter_bytes():
            if not chunk:
                continue
            next_total = total + len(chunk)
            if next_total > max_bytes:
                logger.warning(
                    "proxy 响应超过上限，已中断：url=%s total=%s max=%s",
                    url,
                    next_total,
                    max_bytes,
                )
                break
            total = next_total
            yield chunk

    return StreamingResponse(
        _limited_iter_bytes(),
        status_code=upstream.status_code,
        headers=headers,
        media_type=upstream.headers.get("Content-Type"),
        background=BackgroundTask(_aclose_httpx_stream, upstream, client),
    )

@app.get("/healthz")
def healthz():
    pexels_key = (os.getenv("PEXELS_API_KEY") or "").strip()
    return {
        "ok": True,
        "capabilities": {
            "pexels": {
                # 仅返回布尔值，不暴露 key 内容
                "configured": bool(pexels_key),
            }
        },
    }


if __name__ == "__main__":
    import sys
    import uvicorn

    # 允许在任意工作目录运行：确保可以导入 `backend.common.*`
    repo_root = _find_repo_root(Path(__file__).resolve())
    if str(repo_root) not in sys.path:
        sys.path.insert(0, str(repo_root))

    from backend.common.logging_utils import build_uvicorn_log_config, apply_logging_config

    log_config = build_uvicorn_log_config()
    apply_logging_config(log_config)

    host = os.environ.get("HOST", "0.0.0.0")
    port = int(os.environ.get("MAIN_API_PORT", "6800"))
    uvicorn.run(app, host=host, port=port, log_config=log_config)
