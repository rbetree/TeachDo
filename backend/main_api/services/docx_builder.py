import io
import logging
import re
from typing import Any

from backend.main_api.models.schemas import LessonPlan, LessonStyle
from backend.main_api.utils.common import _normalize_lesson_template_id

logger = logging.getLogger(__name__)

# 教案导出（docx）模板列表
LESSON_DOCX_TEMPLATES: list[dict[str, str]] = [
    {"id": "lesson_simple", "name": "简洁版", "description": "分节标题 + 列表"},
    {"id": "lesson_table", "name": "表格版", "description": "流程表格布局"},
    {"id": "lesson_jnu_form", "name": "教案表单（字段）", "description": "授课题目/授课类型/教学内容/手段与方法/作业/参考资料"},
]


def _lesson_safe_export_filename(title: str) -> str:
    """
    构造可用于 Content-Disposition 的 docx 文件名。
    """
    safe = (title or "").strip().replace("\\", "_").replace("/", "_")
    safe = safe.replace("\r", "_").replace("\n", "_").replace("\t", "_")
    base = safe or "lesson_plan"
    if base.lower().endswith(".docx"):
        return base
    return f"{base}.docx"


def _build_lesson_docx_bytes(*, plan: LessonPlan, style: LessonStyle, language: str, template_id: str | None = None) -> bytes:
    """
    生成 docx bytes（python-docx）。
    - 支持按 template_id 选择不同版式（类似 PPT 模板选择）
    """
    try:
        from docx import Document
        from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, Cm
    except Exception as exc:  # pragma: no cover - 依赖缺失时给出明确错误
        raise RuntimeError(f"缺少 python-docx 依赖：{exc}") from exc

    lang = (language or "zh").strip().lower()
    want_english = lang in {"en", "english"}

    tpl = _normalize_lesson_template_id(template_id)

    def _safe_float(value: Any, fallback: float) -> float:
        try:
            return float(value)
        except Exception:
            logger.debug("将 %r 转为 float 失败，使用 fallback %s", value, fallback, exc_info=True)
            return fallback

    def _safe_text(value: Any, fallback: str) -> str:
        text = str(value or "").strip()
        return text if text else fallback

    def _set_style_font(doc: Any, style_name: str, *, size_pt: int):
        try:
            st = doc.styles[style_name]
        except Exception:
            logger.debug("获取样式 %s 失败", style_name, exc_info=True)
            return
        st.font.name = style.fontZh
        st.font.size = Pt(int(size_pt))
        try:
            st._element.rPr.rFonts.set(qn("w:eastAsia"), style.fontZh)
        except Exception:
            logger.debug("设置样式 %s 中文字体失败", style_name, exc_info=True)
            pass

    def _set_run_font(run: Any, *, size_pt: int, bold: bool = False):
        run.font.name = style.fontZh
        run.font.size = Pt(int(size_pt))
        run.bold = bool(bold)
        try:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), style.fontZh)
        except Exception:
            logger.debug("设置 run 中文字体失败", exc_info=True)
            pass

    def _add_paragraph(doc: Any, text: str, *, size_pt: int, bold: bool = False, style_name: str | None = None, alignment: Any | None = None):
        p = doc.add_paragraph()
        if style_name:
            try:
                p.style = style_name
            except Exception:
                logger.debug("设置段落样式 %s 失败", style_name, exc_info=True)
                pass
        if alignment is not None:
            try:
                p.alignment = alignment
            except Exception:
                logger.debug("设置段落对齐失败", exc_info=True)
                pass
        run = p.add_run(_safe_text(text, "N/A" if want_english else "—"))
        _set_run_font(run, size_pt=size_pt, bold=bold)
        try:
            p.paragraph_format.line_spacing = float(style.lineSpacing)
        except Exception:
            logger.debug("设置段落行距失败", exc_info=True)
            pass
        return p

    def _set_cell_text(cell: Any, text: str, *, size_pt: int, bold: bool = False, alignment: Any | None = None):
        # cell.text 会重置段落样式，这里用 paragraph/run 精细控制字体
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        p.text = ""
        if alignment is not None:
            try:
                p.alignment = alignment
            except Exception:
                logger.debug("设置单元格对齐失败", exc_info=True)
                pass
        run = p.add_run(_safe_text(text, "N/A" if want_english else "—"))
        _set_run_font(run, size_pt=size_pt, bold=bold)
        try:
            p.paragraph_format.line_spacing = float(style.lineSpacing)
        except Exception:
            logger.debug("设置单元格行距失败", exc_info=True)
            pass
        return p

    def _set_cell_paragraphs(
        cell: Any,
        paragraphs: list[tuple[str, bool]],
        *,
        size_pt: int,
        alignment: Any | None = None,
    ):
        """
        以"多段落"方式写入单元格，便于实现"标签（加粗）+ 内容（多行）"的表单样式。
        - paragraphs: [(text, bold), ...]
        """
        # cell.text 会重置段落，这里借助它清空后再逐段写入
        cell.text = ""
        for idx, (text, bold) in enumerate(paragraphs):
            p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
            p.text = ""
            if alignment is not None:
                try:
                    p.alignment = alignment
                except Exception:
                    logger.debug("设置单元格段落对齐失败", exc_info=True)
                    pass
            raw = str(text or "")
            if raw:
                run = p.add_run(raw)
                _set_run_font(run, size_pt=size_pt, bold=bool(bold))
            try:
                p.paragraph_format.line_spacing = float(style.lineSpacing)
            except Exception:
                logger.debug("设置单元格段落行距失败", exc_info=True)
                pass

    def _shade_cell(cell: Any, *, fill_hex: str):
        """
        给单元格加底色（用于表头/标签列）。
        - 不强依赖，失败则忽略（避免不同 docx 实现差异导致导出失败）
        """
        try:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill_hex)
            tc_pr.append(shd)
        except Exception:
            logger.debug("设置单元格底色失败", exc_info=True)
            return

    def _prepare_doc() -> Any:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(_safe_float(style.marginTopCm, 2.54))
        section.bottom_margin = Cm(_safe_float(style.marginBottomCm, 2.54))
        section.left_margin = Cm(_safe_float(style.marginLeftCm, 2.54))
        section.right_margin = Cm(_safe_float(style.marginRightCm, 2.54))

        # 常用样式字体（尽量覆盖列表/标题，避免导出后字体不一致）
        _set_style_font(doc, "Normal", size_pt=style.bodySizePt)
        _set_style_font(doc, "Title", size_pt=style.titleSizePt)
        _set_style_font(doc, "Heading 1", size_pt=style.h1SizePt)
        _set_style_font(doc, "Heading 2", size_pt=style.h2SizePt)
        _set_style_font(doc, "List Bullet", size_pt=style.bodySizePt)
        _set_style_font(doc, "List Number", size_pt=style.bodySizePt)
        _set_style_font(doc, "Table Grid", size_pt=style.bodySizePt)
        return doc

    def _build_common_sections(doc: Any):
        """
        共用分节内容（目标/材料/作业）。不同模板仅在“元信息/流程”呈现上有差异。
        """
        # Objectives
        _add_paragraph(doc, "Objectives" if want_english else "教学目标", size_pt=style.h1SizePt, bold=True, style_name="Heading 1")
        if plan.objectives:
            for item in plan.objectives:
                _add_paragraph(doc, str(item), size_pt=style.bodySizePt, style_name="List Bullet")
        else:
            _add_paragraph(
                doc,
                "No objectives provided." if want_english else "未提供教学目标。",
                size_pt=style.bodySizePt,
                style_name="List Bullet",
            )

        # Materials
        _add_paragraph(doc, "Materials" if want_english else "教学材料", size_pt=style.h1SizePt, bold=True, style_name="Heading 1")
        if plan.materials:
            for item in plan.materials:
                _add_paragraph(doc, str(item), size_pt=style.bodySizePt, style_name="List Bullet")
        else:
            _add_paragraph(
                doc,
                "No materials provided." if want_english else "未提供教学材料。",
                size_pt=style.bodySizePt,
                style_name="List Bullet",
            )

        # Homework
        _add_paragraph(doc, "Homework" if want_english else "课后作业", size_pt=style.h1SizePt, bold=True, style_name="Heading 1")
        _add_paragraph(
            doc,
            (plan.homework or "").strip() or ("No homework provided." if want_english else "未提供课后作业。"),
            size_pt=style.bodySizePt,
        )

    def _build_simple(doc: Any):
        # Title
        _add_paragraph(
            doc,
            (plan.title or "").strip() or ("Lesson Plan" if want_english else "教案"),
            size_pt=style.titleSizePt,
            bold=True,
            style_name="Title",
        )

        # Meta（单行）
        aud_label = "Audience" if want_english else "受众"
        dur_label = "Duration" if want_english else "时长"
        _add_paragraph(
            doc,
            f"{aud_label}：{(plan.targetAudience or '').strip() or ('Students' if want_english else '中学学生')} | "
            f"{dur_label}：{(plan.duration or '').strip() or ('45 min' if want_english else '45分钟')}",
            size_pt=style.bodySizePt,
        )

        # Procedure（编号列表）
        _add_paragraph(doc, "Procedure" if want_english else "教学流程", size_pt=style.h1SizePt, bold=True, style_name="Heading 1")
        if plan.procedure:
            for item in plan.procedure:
                step = (item.step or "").strip() or ("Untitled Step" if want_english else "未命名步骤")
                duration = (item.duration or "").strip() or ("N/A" if want_english else "未填写时长")
                activity = (item.activity or "").strip() or ("No activity details." if want_english else "未填写活动说明。")
                if want_english:
                    _add_paragraph(doc, f"{step} ({duration})", size_pt=style.bodySizePt, bold=True, style_name="List Number")
                else:
                    _add_paragraph(doc, f"{step}（{duration}）", size_pt=style.bodySizePt, bold=True, style_name="List Number")
                _add_paragraph(doc, activity, size_pt=style.bodySizePt)
        else:
            _add_paragraph(
                doc,
                "No procedure provided." if want_english else "未提供教学流程。",
                size_pt=style.bodySizePt,
                style_name="List Number",
            )

        _build_common_sections(doc)

    def _build_table(doc: Any):
        # Title（居中）
        _add_paragraph(
            doc,
            (plan.title or "").strip() or ("Lesson Plan" if want_english else "教案"),
            size_pt=style.titleSizePt,
            bold=True,
            style_name="Title",
            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )

        # Meta（表格：受众/时长）
        aud_label = "Audience" if want_english else "受众"
        dur_label = "Duration" if want_english else "时长"
        meta = doc.add_table(rows=1, cols=4)
        try:
            meta.style = "Table Grid"
        except Exception:
            logger.debug("设置 meta 表格样式失败", exc_info=True)
            pass
        try:
            meta.alignment = WD_TABLE_ALIGNMENT.CENTER
        except Exception:
            logger.debug("设置 meta 表格对齐失败", exc_info=True)
            pass
        for cell in meta.rows[0].cells:
            try:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            except Exception:
                logger.debug("设置 meta 单元格垂直对齐失败", exc_info=True)
                pass

        meta_labels = [aud_label, (plan.targetAudience or "").strip() or ("Students" if want_english else "中学学生"), dur_label, (plan.duration or "").strip() or ("45 min" if want_english else "45分钟")]
        for idx, cell in enumerate(meta.rows[0].cells):
            is_label = idx % 2 == 0
            if is_label:
                _shade_cell(cell, fill_hex="F3F4F6")  # 灰底
            _set_cell_text(
                cell,
                meta_labels[idx],
                size_pt=style.bodySizePt,
                bold=is_label,
                alignment=WD_PARAGRAPH_ALIGNMENT.CENTER if is_label else WD_PARAGRAPH_ALIGNMENT.LEFT,
            )

        # Procedure（表格：环节/时长/活动）
        _add_paragraph(doc, "Procedure" if want_english else "教学流程", size_pt=style.h1SizePt, bold=True, style_name="Heading 1")
        items = plan.procedure or []
        rows = (len(items) if items else 1) + 1  # + 表头
        table = doc.add_table(rows=rows, cols=3)
        try:
            table.style = "Table Grid"
        except Exception:
            logger.debug("设置 procedure 表格样式失败", exc_info=True)
            pass
        try:
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
        except Exception:
            logger.debug("设置 procedure 表格对齐失败", exc_info=True)
            pass
        try:
            table.autofit = False
        except Exception:
            logger.debug("设置 procedure 表格 autofit 失败", exc_info=True)
            pass

        # 尝试设置列宽（不保证所有 Word 版本完全一致，但可提升可读性）
        try:
            table.columns[0].width = Cm(4)
            table.columns[1].width = Cm(3)
            table.columns[2].width = Cm(10)
        except Exception:
            logger.debug("设置 procedure 表格列宽失败", exc_info=True)
            pass

        headers = ["Step" if want_english else "环节", "Duration" if want_english else "时长", "Activity" if want_english else "活动"]
        header_row = table.rows[0]
        for i, cell in enumerate(header_row.cells):
            _shade_cell(cell, fill_hex="E5E7EB")  # 深一点的灰
            try:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            except Exception:
                logger.debug("设置 procedure 表头垂直对齐失败", exc_info=True)
                pass
            _set_cell_text(cell, headers[i], size_pt=style.bodySizePt, bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)

        if items:
            for idx, item in enumerate(items, start=1):
                step = (item.step or "").strip() or ("Untitled Step" if want_english else "未命名步骤")
                duration = (item.duration or "").strip() or ("N/A" if want_english else "未填写时长")
                activity = (item.activity or "").strip() or ("No activity details." if want_english else "未填写活动说明。")
                row = table.rows[idx]
                for cell in row.cells:
                    try:
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                    except Exception:
                        logger.debug("设置 procedure 单元格垂直对齐失败", exc_info=True)
                        pass
                _set_cell_text(row.cells[0], f"{idx}. {step}", size_pt=style.bodySizePt, bold=True)
                _set_cell_text(row.cells[1], duration, size_pt=style.bodySizePt, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
                _set_cell_text(row.cells[2], activity, size_pt=style.bodySizePt)
        else:
            row = table.rows[1]
            _set_cell_text(row.cells[0], "—", size_pt=style.bodySizePt)
            _set_cell_text(row.cells[1], "—", size_pt=style.bodySizePt)
            _set_cell_text(row.cells[2], "No procedure provided." if want_english else "未提供教学流程。", size_pt=style.bodySizePt)

        _build_common_sections(doc)

    def _build_jnu_form(doc: Any):
        """
        “字段表单”模板（参考用户提供的教案表格）：
        - 授课题目、授课类型、授课时间（日期/周次节次）
        - 教学内容（基本内容/重点/难点）
        - 教学手段与方法、作业、参考资料
        """

        # 表头/字段标签（根据语言做最小翻译；该模板主要面向中文）
        if want_english:
            topic_label = "Teaching Topic (chapter/theme):"
            lesson_type_label = "Lesson Type"
            lesson_time_label = "Lesson Time"
            date_placeholder = "YYYY  MM  DD"
            period_placeholder = "Week __  Day __  Period __"
            content_label = "Teaching Content (basics / key / difficult points):"
            methods_label = "Teaching Tools and Methods:"
            homework_label = "Questions / Discussion / Homework:"
            refs_label = "References (books, papers, etc.):"
            basics_label = "Basics:"
            key_label = "Key Points:"
            diff_label = "Difficult Points:"
            tools_label = "Tools:"
            method_label = "Methods:"
        else:
            topic_label = "授课题目（教学章节或主题）："
            lesson_type_label = "授课类型"
            lesson_time_label = "授课时间"
            date_placeholder = "年   月   日"
            period_placeholder = "第   周星期     第   节"
            content_label = "教学内容（包括基本内容、重点、难点三部分）："
            methods_label = "教学手段与方法："
            homework_label = "思考题、讨论题或作业："
            refs_label = "参考资料（包括辅助教材、参考书、文献等）："
            basics_label = "基本内容："
            key_label = "重点："
            diff_label = "难点："
            tools_label = "教学手段："
            method_label = "教学方法："

        def _infer_lesson_type() -> str:
            """
            尝试根据流程/活动粗略推断授课类型。
            - 仅用于填表默认值，避免空白
            """
            haystack = " ".join(
                [
                    f"{(p.step or '').strip()} {(p.activity or '').strip()}".strip()
                    for p in (plan.procedure or [])
                ]
            )
            if want_english:
                if "实验" in haystack or "experiment" in haystack.lower():
                    return "Lab"
                if "讨论" in haystack or "discussion" in haystack.lower():
                    return "Discussion"
                if "实践" in haystack or "实习" in haystack or "见习" in haystack:
                    return "Practice"
                return "Lecture"
            else:
                if "实验" in haystack:
                    return "实验课"
                if "讨论" in haystack:
                    return "讨论课"
                if "实践" in haystack or "实习" in haystack or "见习" in haystack:
                    return "实践课"
                return "理论课"

        def _build_basic_content_lines() -> list[str]:
            if not plan.procedure:
                return ["—" if not want_english else "N/A"]
            lines: list[str] = []
            for idx, item in enumerate(plan.procedure, start=1):
                step = (item.step or "").strip() or ("Untitled Step" if want_english else "未命名步骤")
                duration = (item.duration or "").strip()
                activity = (item.activity or "").strip()
                if duration and activity:
                    lines.append(f"{idx}. {step} ({duration}) {activity}" if want_english else f"{idx}. {step}（{duration}）{activity}")
                elif duration:
                    lines.append(f"{idx}. {step} ({duration})" if want_english else f"{idx}. {step}（{duration}）")
                elif activity:
                    lines.append(f"{idx}. {step}: {activity}" if want_english else f"{idx}. {step}：{activity}")
                else:
                    lines.append(f"{idx}. {step}")
            return lines

        def _build_methods() -> tuple[str, str]:
            # 手段：优先用 materials 列表
            tools = "、".join([str(x).strip() for x in (plan.materials or []) if str(x).strip()]) if not want_english else ", ".join([str(x).strip() for x in (plan.materials or []) if str(x).strip()])
            tools = tools or ("—" if not want_english else "N/A")

            # 方法：用简单规则从流程/活动中抽取关键词
            haystack = " ".join(
                [
                    f"{(p.step or '').strip()} {(p.activity or '').strip()}".strip()
                    for p in (plan.procedure or [])
                ]
            )
            method_candidates: list[str] = []
            if want_english:
                method_candidates.append("Lecture")
                if "讨论" in haystack or "discussion" in haystack.lower():
                    method_candidates.append("Discussion")
                if "示范" in haystack or "demo" in haystack.lower() or "demonstr" in haystack.lower():
                    method_candidates.append("Demonstration")
                if "练习" in haystack or "exercise" in haystack.lower() or "practice" in haystack.lower():
                    method_candidates.append("Practice")
                if "小组" in haystack or "合作" in haystack or "group" in haystack.lower():
                    method_candidates.append("Group work")
                if "提问" in haystack or "问答" in haystack or "q&a" in haystack.lower() or "question" in haystack.lower():
                    method_candidates.append("Q&A")
                # 去重保持顺序
                uniq: list[str] = []
                for m in method_candidates:
                    if m not in uniq:
                        uniq.append(m)
                methods = ", ".join(uniq) if uniq else "N/A"
            else:
                method_candidates.append("讲授")
                if "讨论" in haystack:
                    method_candidates.append("讨论")
                if "示范" in haystack:
                    method_candidates.append("示范")
                if "练习" in haystack or "训练" in haystack:
                    method_candidates.append("练习")
                if "小组" in haystack or "合作" in haystack:
                    method_candidates.append("小组合作")
                if "提问" in haystack or "问答" in haystack:
                    method_candidates.append("提问")
                uniq = []
                for m in method_candidates:
                    if m not in uniq:
                        uniq.append(m)
                methods = "、".join(uniq) if uniq else "—"

            return tools, methods

        def _build_refs_lines() -> list[str]:
            # 参考资料：优先从 materials 中挑选“看起来像参考书/教材/文献”的项；否则留空
            keywords = ("教材", "课本", "参考", "文献", "论文", "书", "ISBN")
            refs = [str(x).strip() for x in (plan.materials or []) if str(x).strip() and any(k in str(x) for k in keywords)]
            if not refs:
                return ["—" if not want_english else "N/A"]
            return [f"- {r}" for r in refs] if not want_english else [f"- {r}" for r in refs]

        lesson_type_value = _infer_lesson_type()

        # 7x3 表格，与用户示例一致：左侧（题目）跨 3 行，授课时间标签跨 2 行，后续 4 行各自跨 3 列
        table = doc.add_table(rows=7, cols=3)
        try:
            table.style = "Table Grid"
        except Exception:
            logger.debug("设置 JNU 表格样式失败", exc_info=True)
            pass
        try:
            table.autofit = False
        except Exception:
            logger.debug("设置 JNU 表格 autofit 失败", exc_info=True)
            pass

        # 尝试设置列宽：A4 扣边距后约 15.9cm，可按需微调
        try:
            table.columns[0].width = Cm(9.0)
            table.columns[1].width = Cm(2.6)
            table.columns[2].width = Cm(4.3)
        except Exception:
            logger.debug("设置 JNU 表格列宽失败", exc_info=True)
            pass

        # 合并单元格（python-docx 用矩形区域合并）
        topic_cell = table.cell(0, 0).merge(table.cell(2, 0))
        time_label_cell = table.cell(1, 1).merge(table.cell(2, 1))
        for r in range(3, 7):
            table.cell(r, 0).merge(table.cell(r, 2))

        # 基础对齐
        for r in range(7):
            for c in range(3):
                cell = table.cell(r, c)
                try:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                except Exception:
                    logger.debug("设置 JNU 单元格垂直对齐失败", exc_info=True)
                    pass

        # Row 0-2: 授课题目 / 类型 / 时间
        _set_cell_paragraphs(
            topic_cell,
            [
                (topic_label, True),
                (_safe_text(plan.title, "Lesson Plan" if want_english else "教案"), False),
            ],
            size_pt=style.bodySizePt,
        )

        _shade_cell(table.cell(0, 1), fill_hex="F3F4F6")
        _set_cell_text(table.cell(0, 1), lesson_type_label, size_pt=style.bodySizePt, bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
        _set_cell_text(table.cell(0, 2), lesson_type_value, size_pt=style.bodySizePt)

        _shade_cell(time_label_cell, fill_hex="F3F4F6")
        _set_cell_text(time_label_cell, lesson_time_label, size_pt=style.bodySizePt, bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
        _set_cell_text(table.cell(1, 2), date_placeholder, size_pt=style.bodySizePt)
        _set_cell_text(table.cell(2, 2), period_placeholder, size_pt=style.bodySizePt)

        # Row 3: 教学内容（基本/重点/难点）
        basic_lines = _build_basic_content_lines()

        obj_list = [str(x).strip() for x in (plan.objectives or []) if str(x).strip()]
        key_objs: list[str] = []
        diff_text = "N/A" if want_english else "—"
        if obj_list:
            last = obj_list[-1]
            is_diff = bool(re.match(r"^difficulty\s*[:：]", last, re.IGNORECASE)) if want_english else bool(re.match(r"^难点\s*[:：]", last))
            if is_diff:
                key_objs = obj_list[:-1]
                diff_text = re.sub(r"^difficulty\s*[:：]\s*", "", last, flags=re.IGNORECASE) if want_english else re.sub(r"^难点\s*[:：]\s*", "", last)
            elif len(obj_list) >= 2:
                key_objs = obj_list[:-1]
                diff_text = last
            else:
                key_objs = obj_list

        key_lines = [f"- {x}" for x in key_objs if str(x).strip()] or ["—" if not want_english else "N/A"]
        diff_lines = [(diff_text or ("N/A" if want_english else "—"), False)]
        content_cell = table.cell(3, 0)
        paragraphs: list[tuple[str, bool]] = [(content_label, True), (basics_label, True)]
        paragraphs.extend([(line, False) for line in basic_lines])
        paragraphs.append(("", False))
        paragraphs.append((key_label, True))
        paragraphs.extend([(line, False) for line in key_lines])
        paragraphs.append(("", False))
        paragraphs.append((diff_label, True))
        paragraphs.extend(diff_lines)
        _set_cell_paragraphs(content_cell, paragraphs, size_pt=style.bodySizePt)

        # Row 4: 教学手段与方法
        tools_value, methods_value = _build_methods()
        methods_cell = table.cell(4, 0)
        _set_cell_paragraphs(
            methods_cell,
            [
                (methods_label, True),
                (f"{tools_label} {tools_value}", False),
                (f"{method_label} {methods_value}", False),
            ],
            size_pt=style.bodySizePt,
        )

        # Row 5: 作业/讨论题
        _set_cell_paragraphs(
            table.cell(5, 0),
            [
                (homework_label, True),
                (_safe_text(plan.homework, "—" if not want_english else "N/A"), False),
            ],
            size_pt=style.bodySizePt,
        )

        # Row 6: 参考资料
        ref_lines = _build_refs_lines()
        _set_cell_paragraphs(
            table.cell(6, 0),
            [(refs_label, True)] + [(line, False) for line in ref_lines],
            size_pt=style.bodySizePt,
        )

    doc = _prepare_doc()
    if tpl == "lesson_simple":
        _build_simple(doc)
    elif tpl == "lesson_table":
        _build_table(doc)
    elif tpl == "lesson_jnu_form":
        _build_jnu_form(doc)
    else:
        raise ValueError(f"未知教案模板：{tpl}")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
