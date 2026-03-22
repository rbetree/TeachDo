from __future__ import annotations

from typing import Any, Dict

import httpx
import pytest
from fastapi import FastAPI, Request
from fastapi.testclient import TestClient


@pytest.fixture()
def main_api_client():
    from backend.main_api.main import app

    return TestClient(app)


@pytest.fixture()
def personaldb_stub() -> Dict[str, Any]:
    """
    personaldb stub：提供 /healthz、/search、/files/{user_id}/{file_id}/content，供 main_api 的 Lesson KB 增强单测使用。
    """
    app = FastAPI()
    state: Dict[str, Any] = {"search_payload": None, "content_calls": []}

    @app.get("/healthz")
    async def healthz():
        return {"ok": True}

    @app.post("/search")
    async def search(request: Request):
        payload = await request.json()
        state["search_payload"] = payload
        return {
            "documents": [["KB chunk 1", "KB chunk 2"]],
            "metadatas": [
                [
                    {
                        "file_id": "upload:test:fid0",
                        "file_name": "a.txt",
                        "folder_id": 0,
                    }
                ]
            ],
            "distances": [[0.1]],
        }

    @app.get("/files/{user_id}/{file_id}/content")
    async def file_content(user_id: str, file_id: str):
        state["content_calls"].append({"user_id": user_id, "file_id": file_id})
        return {
            "user_id": user_id,
            "file_id": file_id,
            "file_name": "lesson.md",
            "file_type": "md",
            "file_size": 10,
            "content": "FULL TEXT FROM gen: OUTPUT",
        }

    return {"app": app, "state": state}


def test_lesson_plan_llm_prompt_includes_course_outputs_and_search_only_uses_rag_ids(
    main_api_client, personaldb_stub, monkeypatch
):
    monkeypatch.setenv("PERSONAL_DB", "http://personaldb.test")
    monkeypatch.setenv("OUTLINE_TYPE", "openai")
    monkeypatch.setenv("OUTLINE_MODEL", "mock-model")
    monkeypatch.setenv("OUTLINE_API_KEY", "sk-test")
    monkeypatch.setenv("OUTLINE_BASE_URL", "http://llm.test/v1")

    import backend.main_api.main as main_api

    stub_app = personaldb_stub["app"]
    stub_state = personaldb_stub["state"]

    real_async_client = httpx.AsyncClient

    def _async_client_factory(*_args, **kwargs):
        timeout = kwargs.get("timeout")
        return real_async_client(
            transport=httpx.ASGITransport(app=stub_app),
            base_url="http://personaldb.test",
            timeout=timeout,
            trust_env=False,
        )

    monkeypatch.setattr(main_api.httpx, "AsyncClient", _async_client_factory)

    seen: Dict[str, Any] = {"system_prompts": []}

    async def _fake_iter_assistant_text_chunks(*, model, api_key, base_url, messages, temperature=0.6):
        # 每个 section 都会调用一次，这里记录 system prompt 用于断言
        seen["system_prompts"].append((messages[0] or {}).get("content", ""))
        user_prompt = (messages[-1] or {}).get("content", "")
        if "targetAudience" in user_prompt:
            yield '{"targetAudience":"中学学生","duration":"45分钟"}'
        elif '"objectives"' in user_prompt:
            yield '{"objectives":["目标1","目标2"]}'
        elif '"materials"' in user_prompt:
            yield '{"materials":["课件/PPT","白板","教材：示例教材"]}'
        elif '"procedure"' in user_prompt:
            yield '{"procedure":[{"step":"导入","duration":"5分钟","activity":"复习并引入"},{"step":"讲解","duration":"15分钟","activity":"讲解核心概念"}]}'
        elif '"homework"' in user_prompt:
            yield '{"homework":"完成课后练习 1~3 题。"}'
        else:
            yield "{}"

    monkeypatch.setattr(main_api, "iter_assistant_text_chunks", _fake_iter_assistant_text_chunks)

    payload = {
        "title": "细胞结构与功能",
        "outlineContent": "# 细胞\n- 细胞膜\n- 细胞核\n",
        "language": "zh",
        "user_id": "default_user",
        "kb_file_ids": ["upload:test:fid0", "gen:test:fid1"],
    }

    with main_api_client.stream("POST", "/tools/lesson_plan", json=payload) as resp:
        assert resp.status_code == 200
        body = b"".join(resp.iter_bytes()).decode("utf-8", errors="ignore")

    assert "[DONE]" in body

    assert stub_state["search_payload"]["userId"] == "default_user"
    assert stub_state["search_payload"]["fileIds"] == ["upload:test:fid0"]
    assert stub_state["content_calls"] == [{"user_id": "default_user", "file_id": "gen:test:fid1"}]

    system_prompt = "\n\n".join([p for p in seen["system_prompts"] if p]).strip()
    assert "课程产出（全文，不经检索）" in system_prompt
    assert "FULL TEXT FROM gen: OUTPUT" in system_prompt
    assert "参考资料检索片段（RAG）" in system_prompt
    assert "KB chunk 1" in system_prompt


def test_lesson_plan_stream_returns_sections_and_done_without_llm_config(main_api_client, monkeypatch):
    # 该用例验证“无 LLM 配置时”的兜底输出，避免被本机 .env 或其他用例的模型配置干扰
    for k in [
        "LESSON_TYPE",
        "LESSON_MODEL",
        "LESSON_API_KEY",
        "LESSON_BASE_URL",
        "OUTLINE_TYPE",
        "OUTLINE_MODEL",
        "OUTLINE_API_KEY",
        "OUTLINE_BASE_URL",
    ]:
        monkeypatch.delenv(k, raising=False)

    payload = {
        "title": "三角形的基本性质",
        "subject": "数学",
        "description": "八年级上册",
        "objectives": "掌握内角和定理；能解决基础题。",
        "outlineContent": "# 三角形\n## 内角和\n- 定理\n- 例题\n",
        "language": "zh",
        "sessionId": "material-1",
        "kb_file_ids": [],
    }

    with main_api_client.stream("POST", "/tools/lesson_plan", json=payload) as resp:
        assert resp.status_code == 200
        body = b"".join(resp.iter_bytes()).decode("utf-8", errors="ignore")

    assert "[DONE]" in body
    assert '"type":"section"' in body
    assert '"section":"objectives"' in body
    assert '"section":"materials"' in body
    assert '"section":"procedure"' in body
    assert '"section":"homework"' in body
    assert '"type":"final"' in body


def test_lesson_plan_fallback_respects_jnu_form_template(main_api_client, monkeypatch):
    # 该用例验证“无 LLM 配置时”的兜底输出，避免被本机 .env 或其他用例的模型配置干扰
    for k in [
        "LESSON_TYPE",
        "LESSON_MODEL",
        "LESSON_API_KEY",
        "LESSON_BASE_URL",
        "OUTLINE_TYPE",
        "OUTLINE_MODEL",
        "OUTLINE_API_KEY",
        "OUTLINE_BASE_URL",
    ]:
        monkeypatch.delenv(k, raising=False)

    payload = {
        "title": "三角形的基本性质",
        "subject": "数学",
        "description": "八年级上册",
        "objectives": "理解内角和定理；能解决基础题。",
        "outlineContent": "# 三角形\n## 内角和\n- 定理\n- 例题\n",
        "language": "zh",
        "sessionId": "material-1",
        "kb_file_ids": [],
        "templateId": "lesson_jnu_form",
    }

    with main_api_client.stream("POST", "/tools/lesson_plan", json=payload) as resp:
        assert resp.status_code == 200
        body = b"".join(resp.iter_bytes()).decode("utf-8", errors="ignore")

    import json

    final_plan = None
    for event in body.split("\n\n"):
        for line in event.splitlines():
            if not line.startswith("data:"):
                continue
            data = line[len("data:") :].strip()
            if not data or data == "[DONE]" or data.startswith(":"):
                continue
            try:
                obj = json.loads(data)
            except Exception:
                continue
            if obj.get("type") == "final":
                final_plan = obj.get("data")
                break
        if final_plan:
            break

    assert final_plan is not None
    objectives = final_plan.get("objectives") or []
    assert objectives
    assert str(objectives[-1]).startswith("难点：")
    materials = final_plan.get("materials") or []
    assert any("教材" in str(m) or "参考" in str(m) for m in materials)


def test_lesson_plan_stream_enriches_kb_search_when_personaldb_available(
    main_api_client, personaldb_stub, monkeypatch
):
    monkeypatch.setenv("PERSONAL_DB", "http://personaldb.test")
    # 强制走 fallback，避免本机 .env 触发真实 LLM 调用导致用例超时
    for k in [
        "LESSON_TYPE",
        "LESSON_MODEL",
        "LESSON_API_KEY",
        "LESSON_BASE_URL",
        "OUTLINE_TYPE",
        "OUTLINE_MODEL",
        "OUTLINE_API_KEY",
        "OUTLINE_BASE_URL",
    ]:
        monkeypatch.delenv(k, raising=False)

    import backend.main_api.main as main_api

    stub_app = personaldb_stub["app"]
    stub_state = personaldb_stub["state"]

    real_async_client = httpx.AsyncClient

    def _async_client_factory(*_args, **kwargs):
        timeout = kwargs.get("timeout")
        return real_async_client(
            transport=httpx.ASGITransport(app=stub_app),
            base_url="http://personaldb.test",
            timeout=timeout,
            trust_env=False,
        )

    monkeypatch.setattr(main_api.httpx, "AsyncClient", _async_client_factory)

    payload = {
        "title": "细胞结构与功能",
        "outlineContent": "# 细胞\n- 细胞膜\n- 细胞核\n",
        "language": "zh",
        "user_id": "default_user",
        "kb_file_ids": ["upload:test:fid0"],
    }

    with main_api_client.stream("POST", "/tools/lesson_plan", json=payload) as resp:
        assert resp.status_code == 200
        _ = b"".join(resp.iter_bytes())

    assert stub_state["search_payload"]["userId"] == "default_user"
    assert "细胞结构与功能" in stub_state["search_payload"]["query"]
    assert stub_state["search_payload"]["fileIds"] == ["upload:test:fid0"]


def test_lesson_export_docx_returns_docx_file(main_api_client):
    payload = {
        "lessonPlan": {
            "title": "三角形的基本性质",
            "targetAudience": "中学学生",
            "duration": "45分钟",
            "objectives": ["理解内角和定理", "完成基础练习"],
            "materials": ["课件/PPT", "白板"],
            "procedure": [
                {"step": "导入", "duration": "5分钟", "activity": "复习三角形概念并引出问题"},
                {"step": "讲解", "duration": "15分钟", "activity": "讲解内角和定理并示范推导"},
            ],
            "homework": "完成课后练习 1~3 题。",
        },
        "style": {
            "fontZh": "微软雅黑",
            "titleSizePt": 20,
            "h1SizePt": 16,
            "h2SizePt": 14,
            "bodySizePt": 12,
            "lineSpacing": 1.5,
        },
        "language": "zh",
    }

    resp = main_api_client.post("/lesson/export/docx", json=payload)
    assert resp.status_code == 200
    assert resp.headers.get("content-type", "").startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "attachment" in (resp.headers.get("content-disposition") or "")
    assert resp.content[:2] == b"PK"  # docx 是 zip 包
    assert len(resp.content) > 1000


def test_lesson_templates_endpoint_returns_catalog(main_api_client):
    resp = main_api_client.get("/lesson/templates")
    assert resp.status_code == 200
    data = resp.json().get("data") or []
    ids = {item.get("id") for item in data}
    assert "lesson_simple" in ids
    assert "lesson_table" in ids
    assert "lesson_jnu_form" in ids


def test_lesson_export_docx_supports_table_template(main_api_client):
    payload = {
        "lessonPlan": {
            "title": "三角形的基本性质",
            "targetAudience": "中学学生",
            "duration": "45分钟",
            "objectives": ["理解内角和定理", "完成基础练习"],
            "materials": ["课件/PPT", "白板"],
            "procedure": [
                {"step": "导入", "duration": "5分钟", "activity": "复习三角形概念并引出问题"},
                {"step": "讲解", "duration": "15分钟", "activity": "讲解内角和定理并示范推导"},
            ],
            "homework": "完成课后练习 1~3 题。",
        },
        "style": {
            "fontZh": "微软雅黑",
            "titleSizePt": 20,
            "h1SizePt": 16,
            "h2SizePt": 14,
            "bodySizePt": 12,
            "lineSpacing": 1.5,
        },
        "language": "zh",
        "templateId": "lesson_table",
    }

    resp = main_api_client.post("/lesson/export/docx", json=payload)
    assert resp.status_code == 200
    assert resp.content[:2] == b"PK"

    # 校验：表格模板至少包含“元信息表 + 流程表”两张表
    from io import BytesIO

    from docx import Document

    doc = Document(BytesIO(resp.content))
    assert len(doc.tables) >= 2
    assert doc.tables[0].cell(0, 0).text.strip() == "受众"
    assert doc.tables[1].cell(0, 0).text.strip() == "环节"


def test_lesson_export_docx_supports_jnu_form_template(main_api_client):
    payload = {
        "lessonPlan": {
            "title": "三角形的基本性质",
            "targetAudience": "中学学生",
            "duration": "45分钟",
            "objectives": ["理解内角和定理", "完成基础练习"],
            "materials": ["课件/PPT", "白板"],
            "procedure": [
                {"step": "导入", "duration": "5分钟", "activity": "复习三角形概念并引出问题"},
                {"step": "讲解", "duration": "15分钟", "activity": "讲解内角和定理并示范推导"},
            ],
            "homework": "完成课后练习 1~3 题。",
        },
        "style": {
            "fontZh": "微软雅黑",
            "titleSizePt": 20,
            "h1SizePt": 16,
            "h2SizePt": 14,
            "bodySizePt": 12,
            "lineSpacing": 1.5,
        },
        "language": "zh",
        "templateId": "lesson_jnu_form",
    }

    resp = main_api_client.post("/lesson/export/docx", json=payload)
    assert resp.status_code == 200
    assert resp.content[:2] == b"PK"

    from io import BytesIO

    from docx import Document

    doc = Document(BytesIO(resp.content))
    assert len(doc.tables) >= 1
    table = doc.tables[0]
    assert len(table.rows) == 7
    assert len(table.columns) == 3
    assert "授课题目" in table.cell(0, 0).text
    assert "三角形的基本性质" in table.cell(0, 0).text
    assert table.cell(0, 1).text.strip() == "授课类型"
    assert table.cell(1, 1).text.strip() == "授课时间"
    assert "教学内容" in table.cell(3, 0).text
    assert "教学手段与方法" in table.cell(4, 0).text
